from rules.base_rule import BaseRule, RuleResult
from docx.shared import RGBColor

class FontColorRule(BaseRule):
    """字体颜色统一规则 - 将文档中所有文本的颜色统一设置"""

    display_name = "字体颜色统一"
    category = "字体规则"
    
    def __init__(self, config=None):
        default_params = {
            'text_color': '000000',
        }
        super().__init__({**default_params, **(config or {})})
    
    def apply(self, doc_context):
        """
        应用字体颜色规则
        :param doc_context: 文档上下文
        :return: 规则执行结果
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        # 解析颜色配置
        text_color = self.config.get('text_color', '000000')
        # 将十六进制颜色转换为RGB
        if isinstance(text_color, str) and len(text_color) == 6:
            r = int(text_color[0:2], 16)
            g = int(text_color[2:4], 16)
            b = int(text_color[4:6], 16)
            target_color = RGBColor(r, g, b)
        else:
            # 默认黑色
            target_color = RGBColor(0, 0, 0)
        
        # 处理段落文本
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb != target_color:
                    run.font.color.rgb = target_color
                    fixed_count += 1
        
        # 处理表格中的文本
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.color.rgb != target_color:
                                run.font.color.rgb = target_color
                                fixed_count += 1
        
        if fixed_count > 0:
            details.append(f"统一了{fixed_count}个文本的颜色为{text_color}")
        else:
            details.append(f"文档中所有文本颜色已经是{text_color}")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
