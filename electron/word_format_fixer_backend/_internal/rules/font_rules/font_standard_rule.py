from rules.base_rule import BaseRule, RuleResult
from docx.shared import Pt
from docx.oxml.ns import qn

class FontNameRule(BaseRule):
    """字体名称标准化规则 - 设置中英文字体"""

    display_name = "字体名称标准化"
    category = "字体规则"

    def __init__(self, config=None):
        default_params = {
            'chinese_font': '宋体',
            'western_font': 'Arial',
        }
        super().__init__({**default_params, **(config or {})})

    def apply(self, doc_context) -> RuleResult:
        """
        应用字体名称规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.config['western_font']
                run._element.rPr.rFonts.set(
                    qn('w:eastAsia'),
                    self.config['chinese_font']
                )
                fixed_count += 1

        # 处理表格中的文本
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.config['western_font']
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'),
                                self.config['chinese_font']
                            )
                            fixed_count += 1

        details.append(f"设置中文字体: {self.config['chinese_font']}, 西文字体: {self.config['western_font']}")
        details.append(f"标准化了 {fixed_count} 个文本运行的字体")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )


class TitleFontRule(BaseRule):
    """标题字体规则 - 标题使用专用字体"""

    display_name = "标题字体设置"
    category = "字体规则"

    def __init__(self, config=None):
        default_params = {
            'title_font': '黑体',
        }
        super().__init__({**default_params, **(config or {})})

    def apply(self, doc_context) -> RuleResult:
        """
        应用标题字体规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                for run in paragraph.runs:
                    run.font.name = 'Arial'  # 西文字体固定为Arial
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'),
                        self.config['title_font']
                    )
                    fixed_count += 1

        details.append(f"标题使用字体: {self.config['title_font']}")
        details.append(f"标准化了 {fixed_count} 个标题文本运行的字体")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )


class FontSizeRule(BaseRule):
    """字号标准化规则 - 设置标题和正文的字号"""

    display_name = "字号标准化"
    category = "字体规则"

    def __init__(self, config=None):
        default_params = {
            'font_size_body': 12,
            'font_size_title1': 22,
            'font_size_title2': 18,
            'font_size_title3': 16,
            'min_font_size': 10,  # 小于此字号的文本会被标准化
        }
        super().__init__({**default_params, **(config or {})})

    def apply(self, doc_context) -> RuleResult:
        """
        应用字号规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                # 根据标题级别设置字号
                for run in paragraph.runs:
                    if paragraph.style.name == 'Heading 1':
                        run.font.size = Pt(self.config['font_size_title1'])
                    elif paragraph.style.name == 'Heading 2':
                        run.font.size = Pt(self.config['font_size_title2'])
                    elif paragraph.style.name == 'Heading 3':
                        run.font.size = Pt(self.config['font_size_title3'])
                    else:
                        run.font.size = Pt(self.config['font_size_body'])
                    fixed_count += 1
            else:
                # 正文：如果字号太小则标准化
                for run in paragraph.runs:
                    if run.font.size is None or run.font.size.pt < self.config['min_font_size']:
                        run.font.size = Pt(self.config['font_size_body'])
                        fixed_count += 1

        details.append(f"正文字号: {self.config['font_size_body']}pt")
        details.append(f"一级标题字号: {self.config['font_size_title1']}pt")
        details.append(f"二级标题字号: {self.config['font_size_title2']}pt")
        details.append(f"三级标题字号: {self.config['font_size_title3']}pt")
        details.append(f"标准化了 {fixed_count} 个文本运行的字号")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
