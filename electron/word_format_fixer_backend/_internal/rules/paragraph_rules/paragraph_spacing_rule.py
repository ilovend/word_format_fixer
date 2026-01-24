from rules.base_rule import BaseRule, RuleResult
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING

class ParagraphSpacingRule(BaseRule):
    """段落间距统一规则 - 统一文档中所有段落的间距和缩进"""

    display_name = "段落间距统一"
    category = "段落规则"
    
    def __init__(self, config=None):
        default_params = {
            'body_left_indent': 0,  # 正文左缩进（厘米）
            'body_right_indent': 0,  # 正文右缩进（厘米）
            'body_space_before': 0,  # 正文段前间距（厘米）
            'body_space_after': 0.33,  # 正文段后间距（厘米）
            'body_line_spacing': 1.5,  # 正文行间距
            'table_left_indent': 0.2,  # 表格内段落左缩进（厘米）
            'table_right_indent': 0.2,  # 表格内段落右缩进（厘米）
            'table_space_before': 0,  # 表格内段前段落间距（厘米）
            'table_space_after': 0,  # 表格内段后段落间距（厘米）
        }
        super().__init__({**default_params, **(config or {})})
    
    def apply(self, doc_context):
        """
        应用段落间距规则
        :param doc_context: 文档上下文
        :return: 规则执行结果
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        # 处理普通段落
        for paragraph in document.paragraphs:
            # 跳过标题段落
            if not paragraph.style.name.startswith('Heading'):
                # 设置段落格式
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Cm(self.config['body_left_indent'])
                paragraph_format.right_indent = Cm(self.config['body_right_indent'])
                paragraph_format.space_before = Cm(self.config['body_space_before'])
                paragraph_format.space_after = Cm(self.config['body_space_after'])
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph_format.line_spacing = self.config['body_line_spacing']
                fixed_count += 1
        
        # 处理表格中的段落
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Cm(self.config['table_left_indent'])
                        paragraph_format.right_indent = Cm(self.config['table_right_indent'])
                        paragraph_format.space_before = Cm(self.config['table_space_before'])
                        paragraph_format.space_after = Cm(self.config['table_space_after'])
                        fixed_count += 1
        
        if fixed_count > 0:
            details.append(f"统一了{fixed_count}个段落的间距和缩进")
        else:
            details.append("文档中没有需要处理的段落")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
