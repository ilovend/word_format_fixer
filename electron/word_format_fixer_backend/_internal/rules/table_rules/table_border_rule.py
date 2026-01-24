from rules.base_rule import BaseRule, RuleResult
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

class TableBorderRule(BaseRule):
    """表格边框规则 - 为表格添加边框并格式化单元格"""

    display_name = "表格边框和格式"
    category = "表格规则"
    
    def __init__(self, config=None):
        # 搬运原 default_config 中的表格相关配置
        default_params = {
            'chinese_font': '宋体',
            'western_font': 'Arial',
            'font_size_table_header': 14,
            'font_size_table_content': 12,
            'text_color': (0, 0, 0),
            'table_header_bg_color': 'E3E3E3',  # 表头背景色
            'add_table_header_format': True,
            'border_size': 4,
            'border_color': '000000',
        }
        super().__init__({**default_params, **(config or {})})
    
    def apply(self, doc_context) -> RuleResult:
        """
        核心执行逻辑
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        if not document.tables:
            details.append("文档中没有表格，跳过表格边框和格式设置")
            return RuleResult(
                rule_id=self.rule_id,
                success=True,
                fixed_count=fixed_count,
                details=details
            )
        
        details.append(f"开始为表格添加边框和格式（共{len(document.tables)}个）...")
        
        for table_idx, table in enumerate(document.tables):
            details.append(f"处理表格 {table_idx + 1}: {len(table.rows)}行 × {len(table.columns)}列")
            
            # 为表格添加边框
            border_size = self.config.get('border_size', 4)
            border_color = self.config.get('border_color', '000000')
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell, border_size, border_color)
                    fixed_count += 1
            
            # 格式化表格单元格
            self._format_table_cells(table)
            
            details.append(f"完成表格 {table_idx + 1} 的边框和格式设置")
        
        details.append(f"总共处理了 {fixed_count} 个表格单元格")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
    
    def _set_cell_border(self, cell, border_size: int = 4, border_color: str = "000000"):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        borders = ['top', 'left', 'bottom', 'right']
        for border in borders:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), str(border_size))
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), border_color)
            tcPr.append(border_elem)
    
    def _format_table_cells(self, table):
        """格式化表格单元格"""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 设置单元格垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 设置字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.config['western_font']
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'), 
                            self.config['chinese_font']
                        )
                        run.font.color.rgb = RGBColor(*self.config['text_color'])
                        
                        # 根据位置设置字号
                        if i == 0 and self.config['add_table_header_format']:
                            run.font.size = Pt(self.config['font_size_table_header'])
                            run.font.bold = True
                        else:
                            run.font.size = Pt(self.config['font_size_table_content'])
                
                # 设置表头背景色
                if i == 0 and self.config['add_table_header_format']:
                    self._set_cell_background(cell, self.config['table_header_bg_color'])
                
                # 设置单元格边距
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.left_indent = Cm(0.2)
                    cell.paragraphs[0].paragraph_format.right_indent = Cm(0.2)
    
    def _set_cell_background(self, cell, color):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        # 处理十六进制颜色代码
        if isinstance(color, str) and len(color) == 6:
            fill_color = color
        # 处理RGB元组
        elif isinstance(color, tuple) and len(color) == 3:
            fill_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        else:
            fill_color = "E3E3E3"  # 默认灰色
        shading.set(qn('w:fill'), fill_color)
        cell._element.tcPr.append(shading)
