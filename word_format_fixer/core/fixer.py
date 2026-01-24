"""Word文档修复器核心功能"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Dict, List, Tuple


class RobustWordFixer:
    """修复了页面宽度计算问题的Word修复器"""
    
    def __init__(self, config: Dict = None):
        """
        初始化修复器
        
        Args:
            config: 配置字典
        """
        self.default_config = {
            # 字体设置
            'chinese_font': '宋体',
            'western_font': 'Arial',
            'title_font': '黑体',
            
            # 字号设置
            'font_size_body': 12,
            'font_size_title1': 22,
            'font_size_title2': 18,
            'font_size_title3': 16,
            'font_size_table_header': 14,
            'font_size_table_content': 12,
            
            # 颜色设置
            'text_color': (0, 0, 0),
            'table_header_bg_color': (217, 217, 217),  # 表头背景色
            
            # 页面设置 - 更安全的边距
            'page_width_cm': 21.0,  # A4纸宽度
            'page_height_cm': 29.7,  # A4纸高度
            'page_margin_top_cm': 2.54,
            'page_margin_bottom_cm': 2.54,
            'page_margin_left_cm': 2.54,
            'page_margin_right_cm': 2.54,
            
            # 表格设置
            'table_width_percent': 95,  # 表格宽度占页面宽度的百分比
            'table_alignment': WD_TABLE_ALIGNMENT.CENTER,
            'cell_margin_top_cm': 0.1,
            'cell_margin_bottom_cm': 0.1,
            'cell_margin_left_cm': 0.2,
            'cell_margin_right_cm': 0.2,
            'column_widths': [],  # 自定义列宽百分比
            
            # 修复选项
            'fix_table_width': True,
            'adjust_table_layout': True,
            'add_table_header_format': True,
            'center_vertically': False,
            'auto_adjust_columns': True,
        }
        
        self.config = {**self.default_config, **(config or {})}
        self.document = None
        self.input_filepath = None
    
    def load_document(self, filepath: str):
        """加载文档"""
        self.document = Document(filepath)
        self.input_filepath = filepath
        return self.document
    
    def save_document(self, filepath: str = None):
        """保存文档"""
        if not self.document:
            raise ValueError("文档未加载")
        
        if filepath:
            save_path = filepath
        else:
            base, ext = os.path.splitext(self.input_filepath)
            save_path = f"{base}_fixed{ext}"
        
        self.document.save(save_path)
        return save_path
    
    def set_page_layout(self):
        """设置页面布局 - 修复版本"""
        for section in self.document.sections:
            # 确保页面大小已设置
            if section.page_width is None:
                section.page_width = Cm(self.config['page_width_cm'])
            if section.page_height is None:
                section.page_height = Cm(self.config['page_height_cm'])
            
            # 设置边距
            section.top_margin = Cm(self.config.get('page_margin_top_cm', 2.54))
            section.bottom_margin = Cm(self.config.get('page_margin_bottom_cm', 2.54))
            section.left_margin = Cm(self.config.get('page_margin_left_cm', 2.54))
            section.right_margin = Cm(self.config.get('page_margin_right_cm', 2.54))
            
            print("页面设置完成:")
            print(f"  页面宽度: {section.page_width.cm:.1f}cm")
            print(f"  页面高度: {section.page_height.cm:.1f}cm")
            print(f"  左边距: {section.left_margin.cm:.1f}cm")
            print(f"  右边距: {section.right_margin.cm:.1f}cm")
            
            # 计算可用宽度（使用配置的值，而不是section.page_width）
            page_width_cm = self.config['page_width_cm']
            left_margin_cm = self.config['page_margin_left_cm']
            right_margin_cm = self.config['page_margin_right_cm']
            available_width_cm = page_width_cm - left_margin_cm - right_margin_cm
            
            print(f"  可用宽度: {available_width_cm:.1f}cm")
            
            # 保存可用宽度供后续使用
            self.available_width_cm = available_width_cm
    
    def optimize_table_width(self):
        """优化表格宽度 - 支持复杂表格结构"""
        if not self.config['fix_table_width'] or not self.document.tables:
            return
        
        print(f"\n优化表格（共{len(self.document.tables)}个）...")
        
        for table_idx, table in enumerate(self.document.tables):
            print(f"  处理表格 {table_idx + 1}: {len(table.columns)}列")
            
            # 检查是否有合并单元格
            has_merged_cells = self._check_merged_cells(table)
            if has_merged_cells:
                print(f"    注意: 表格包含合并单元格")
            
            # 检查是否有嵌套表格
            has_nested_tables = self._check_nested_tables(table)
            if has_nested_tables:
                print(f"    注意: 表格包含嵌套表格")
            
            # 计算表格宽度
            table_width_cm = self.available_width_cm * self.config['table_width_percent'] / 100
            table.width = Cm(table_width_cm)
            table.alignment = self.config['table_alignment']
            
            print(f"    表格宽度: {table_width_cm:.1f}cm ({self.config['table_width_percent']}%)")
            
            # 处理列宽
            col_count = len(table.columns)
            
            if self.config['column_widths'] and len(self.config['column_widths']) == col_count:
                # 使用自定义列宽
                total_percent = sum(self.config['column_widths'])
                for i, col in enumerate(table.columns):
                    col_width_cm = (table_width_cm * 
                                   self.config['column_widths'][i] / 
                                   total_percent)
                    col.width = Cm(col_width_cm)
                    print(f"    列{i+1}宽度: {col_width_cm:.1f}cm "                           f"({self.config['column_widths'][i]}%)")
            elif self.config['auto_adjust_columns']:
                # 自动调整列宽
                self._auto_adjust_column_widths(table, table_width_cm)
            else:
                # 平均分配列宽
                col_width_cm = table_width_cm / col_count
                for col in table.columns:
                    col.width = Cm(col_width_cm)
                print(f"    平均列宽: {col_width_cm:.1f}cm")
            
            # 处理嵌套表格
            if has_nested_tables:
                self._process_nested_tables(table)
    
    def _check_merged_cells(self, table):
        """检查表格是否有合并单元格"""
        for row in table.rows:
            for cell in row.cells:
                # 检查是否有vMerge或hMerge属性（使用本地名称查询）
                tcPr = cell._tc.get_or_add_tcPr()
                for child in tcPr.iter():
                    local_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if local_name in ['vMerge', 'hMerge']:
                        return True
        return False
    
    def _check_nested_tables(self, table):
        """检查表格是否有嵌套表格"""
        for row in table.rows:
            for cell in row.cells:
                if len(cell.tables) > 0:
                    return True
        return False
    
    def _process_nested_tables(self, table):
        """处理嵌套表格"""
        for row in table.rows:
            for cell in row.cells:
                for nested_table in cell.tables:
                    # 为嵌套表格设置较小的宽度
                    nested_table_width_cm = cell.width.cm * 0.9 if cell.width else self.available_width_cm * 0.7
                    nested_table.width = Cm(nested_table_width_cm)
                    nested_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 为嵌套表格添加边框
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            self._set_cell_border(nested_cell, border_size=2)
                    
                    print(f"    处理了一个嵌套表格，宽度: {nested_table_width_cm:.1f}cm")
    
    def _auto_adjust_column_widths(self, table, table_width_cm):
        """自动调整列宽 - 支持复杂表格"""
        col_count = len(table.columns)
        max_lengths = [0] * col_count
        
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                # 处理合并单元格的情况
                # 检查单元格是否跨列
                tcPr = cell._tc.get_or_add_tcPr()
                # 使用本地名称查询替代带有命名空间前缀的XPath查询
                gridSpan = None
                for child in tcPr.iter():
                    local_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if local_name == 'gridSpan':
                        gridSpan = child
                        break
                span = 1
                if gridSpan is not None:
                    # 查找val属性
                    val = None
                    for attr in gridSpan.attrib:
                        local_attr_name = attr.split('}')[-1] if '}' in attr else attr
                        if local_attr_name == 'val':
                            val = gridSpan.attrib[attr]
                            break
                    if val:
                        span = int(val)
                
                text = cell.text.strip()
                if text:
                    # 估算文本长度（中文占2个字符宽度）
                    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
                    english_chars = len(text) - chinese_chars
                    length = (english_chars + chinese_chars * 2) / span  # 平均分配到跨列
                    
                    # 更新跨列的最大长度
                    for k in range(j, min(j + span, col_count)):
                        max_lengths[k] = max(max_lengths[k], length)
        
        total_length = sum(max_lengths) if sum(max_lengths) > 0 else 1
        
        for j, col in enumerate(table.columns):
            width_ratio = max_lengths[j] / total_length
            col_width_cm = table_width_cm * width_ratio
            col.width = Cm(col_width_cm)
            print(f"    列{j+1}宽度: {col_width_cm:.1f}cm (比例: {width_ratio:.2f})")
    
    def format_table_cells(self):
        """格式化表格单元格"""
        for table in self.document.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # 设置单元格垂直居中
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # 设置字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.config['western_font']
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'), 
                                self.config['chinese_font']
                            )
                            run.font.color.rgb = RGBColor(*self.config['text_color'])
                            
                            # 根据位置设置字号
                            if i == 0 and self.config['add_table_header_format']:
                                run.font.size = Pt(self.config['font_size_table_header'])
                                run.font.bold = True
                            else:
                                run.font.size = Pt(self.config['font_size_table_content'])
                    
                    # 设置表头背景色
                    if i == 0 and self.config['add_table_header_format']:
                        self._set_cell_background(cell, self.config['table_header_bg_color'])
                    
                    # 设置单元格边距
                    cell.paragraphs[0].paragraph_format.left_indent = Cm(0.2)
                    cell.paragraphs[0].paragraph_format.right_indent = Cm(0.2)
    
    def _set_cell_background(self, cell, color_rgb: Tuple[int, int, int]):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
        cell._element.tcPr.append(shading)
    
    def add_table_borders(self):
        """为表格添加边框"""
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell)
    
    def _set_cell_border(self, cell, border_size: int = 4, border_color: str = "000000"):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        borders = ['top', 'left', 'bottom', 'right']
        for border in borders:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), str(border_size))
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), border_color)
            tcPr.append(border_elem)
    
    def fix_fonts_and_colors(self):
        """修复字体和颜色"""
        for paragraph in self.document.paragraphs:
            # 设置标题格式
            if paragraph.style.name.startswith('Heading'):
                # 只让大标题（Heading 1）居中，小标题不居中
                if paragraph.style.name == 'Heading 1':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 标题使用黑体
                for run in paragraph.runs:
                    run.font.name = self.config['western_font']
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 
                        self.config['title_font']
                    )
                    run.font.color.rgb = RGBColor(*self.config['text_color'])
                    
                    # 根据标题级别设置字号
                    if paragraph.style.name == 'Heading 1':
                        run.font.size = Pt(self.config['font_size_title1'])
                        run.font.bold = True
                    elif paragraph.style.name == 'Heading 2':
                        run.font.size = Pt(self.config['font_size_title2'])
                        run.font.bold = True
                    elif paragraph.style.name == 'Heading 3':
                        run.font.size = Pt(self.config['font_size_title3'])
                        run.font.bold = True
                    else:
                        run.font.size = Pt(self.config['font_size_body'])
                        run.font.bold = True
            else:
                # 正文使用宋体
                for run in paragraph.runs:
                    run.font.name = self.config['western_font']
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 
                        self.config['chinese_font']
                    )
                    run.font.color.rgb = RGBColor(*self.config['text_color'])
                    
                    # 设置合适的字号
                    if run.font.size is None or run.font.size.pt < 10:
                        run.font.size = Pt(self.config['font_size_body'])
    
    def detect_numbering_patterns(self):
        """检测文档中的编号模式"""
        patterns = {
            'arabic_with_dot': r'^\s*(\d+)\.\s+',  # 1. 项目
            'arabic_with_paren': r'^\s*(\d+)\)\s+',  # 1) 项目
            'chinese_with_dot': r'^\s*([一二三四五六七八九十]+)\.\s+',  # 一. 项目
            'chinese_with_bracket': r'^\s*([一二三四五六七八九十]+)、\s+',  # 一、项目
            'chinese_with_paren': r'^\s*\(([一二三四五六七八九十]+)\)\s+',  # (一) 项目
            'lower_alpha_with_dot': r'^\s*([a-z])\.\s+',  # a. 项目
            'upper_alpha_with_dot': r'^\s*([A-Z])\.\s+',  # A. 项目
            'lower_roman_with_dot': r'^\s*([ivxlcdm]+)\.\s+',  # i. 项目
            'upper_roman_with_dot': r'^\s*([IVXLCDM]+)\.\s+',  # I. 项目
        }
        return patterns
    
    def fix_numbering_lists(self):
        """修复编号列表和项目符号"""
        print("修复编号列表和项目符号...")
        
        patterns = self.detect_numbering_patterns()
        fixed_count = 0
        
        for paragraph in self.document.paragraphs:
            original_text = paragraph.text.strip()
            
            if not original_text or paragraph.style.name.startswith('Heading'):
                continue
            
            # 处理不需要的项目符号（如·）
            bullet_patterns = [
                r'^\s*·\s+',  # 中文点号
                r'^\s*\*\s+',  # 星号
                r'^\s*-\s+',   # 连字符
                r'^\s+•\s+',   # 实心圆点
            ]
            
            bullet_matched = False
            for bullet_pattern in bullet_patterns:
                bullet_match = re.match(bullet_pattern, original_text)
                if bullet_match:
                    bullet_matched = True
                    # 提取内容（去掉项目符号）
                    content = original_text[len(bullet_match.group()):].strip()
                    
                    # 清除原有内容
                    paragraph.clear()
                    
                    # 添加内容
                    run = paragraph.add_run(content)
                    run.font.name = self.config['western_font']
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 
                        self.config['chinese_font']
                    )
                    run.font.color.rgb = RGBColor(*self.config['text_color'])
                    run.font.size = Pt(self.config['font_size_body'])
                    
                    # 设置为无序列表样式
                    if 'List Paragraph' in [style.name for style in self.document.styles]:
                        paragraph.style = self.document.styles['List Paragraph']
                    
                    # 设置缩进
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Cm(1.27)
                    paragraph_format.first_line_indent = Cm(-0.64)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph_format.line_spacing = 1.5
                    
                    fixed_count += 1
                    break
            
            if bullet_matched:
                continue
            
            # 检查各种编号模式
            for pattern_name, pattern in patterns.items():
                match = re.match(pattern, original_text)
                if match:
                    self._format_numbered_paragraph(paragraph, pattern_name, match, original_text)
                    fixed_count += 1
                    break
        
        print(f"  修复了 {fixed_count} 个编号或项目符号段落")
        return fixed_count
    
    def _format_numbered_paragraph(self, paragraph, pattern_type: str, match, original_text: str):
        """格式化编号段落"""
        # 提取编号和内容
        number_part = match.group()
        content = original_text[len(number_part):].strip()
        
        # 清除原有内容
        paragraph.clear()
        
        # 添加内容
        run = paragraph.add_run(content)
        run.font.name = self.config['western_font']
        run._element.rPr.rFonts.set(
            qn('w:eastAsia'), 
            self.config['chinese_font']
        )
        run.font.color.rgb = RGBColor(*self.config['text_color'])
        run.font.size = Pt(self.config['font_size_body'])
        
        # 设置段落格式
        if 'List Paragraph' in [style.name for style in self.document.styles]:
            paragraph.style = self.document.styles['List Paragraph']
        else:
            paragraph.style = self.document.styles['Normal']
        
        # 设置缩进
        paragraph_format = paragraph.paragraph_format
        
        # 根据编号类型设置不同缩进
        if pattern_type.startswith('arabic') or pattern_type.startswith('chinese'):
            # 主要列表项
            paragraph_format.left_indent = Cm(1.27)
            paragraph_format.first_line_indent = Cm(-0.64)
        elif pattern_type.startswith('lower_') or pattern_type.startswith('upper_'):
            # 次级列表项
            paragraph_format.left_indent = Cm(2.54)
            paragraph_format.first_line_indent = Cm(-0.64)
        else:
            # 默认缩进
            paragraph_format.left_indent = Cm(1.27)
            paragraph_format.first_line_indent = Cm(-0.64)
        
        # 设置行距
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
    
    def fix_all(self, filepath: str, output_path: str = None):
        """执行所有修复"""
        print("=" * 60)
        print(f"开始修复文档格式: {filepath}")
        print("=" * 60)
        
        try:
            # 加载文档
            self.load_document(filepath)
            
            # 执行修复步骤
            print("\n1. 设置页面布局...")
            self.set_page_layout()
            
            print("\n2. 修复字体和颜色...")
            self.fix_fonts_and_colors()
            
            print("\n3. 修复编号列表...")
            self.fix_numbering_lists()
            
            if self.document.tables:
                print(f"\n4. 优化表格（共{len(self.document.tables)}个）...")
                self.optimize_table_width()
                self.format_table_cells()
                self.add_table_borders()
            else:
                print("\n4. 文档中没有表格，跳过表格优化")
            
            # 保存文档
            save_path = self.save_document(output_path)
            
            print("\n" + "=" * 60)
            print("文档修复完成！")
            print(f"输出文件: {save_path}")
            print("=" * 60)
            
            return save_path
            
        except Exception as e:
            print(f"\n发生错误: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def fix_batch(self, filepaths: list, output_dir: str = None):
        """批量修复多个文档
        
        Args:
            filepaths: 文件路径列表
            output_dir: 输出目录，None表示在原文件目录保存
            
        Returns:
            修复结果字典，键为输入文件路径，值为输出文件路径或None（失败）
        """
        print("=" * 60)
        print(f"开始批量修复，共{len(filepaths)}个文件")
        print("=" * 60)
        
        results = {}
        total_success = 0
        total_failed = 0
        
        for i, filepath in enumerate(filepaths, 1):
            print(f"\n{'-' * 60}")
            print(f"处理文件 {i}/{len(filepaths)}: {filepath}")
            print(f"{'-' * 60}")
            
            try:
                # 生成输出路径
                if output_dir:
                    filename = os.path.basename(filepath)
                    base, ext = os.path.splitext(filename)
                    output_path = os.path.join(output_dir, f"{base}_fixed{ext}")
                else:
                    output_path = None
                
                # 执行修复
                result = self.fix_all(filepath, output_path)
                
                if result:
                    results[filepath] = result
                    total_success += 1
                    print(f"✓ 修复成功: {result}")
                else:
                    results[filepath] = None
                    total_failed += 1
                    print(f"✗ 修复失败")
                    
            except Exception as e:
                results[filepath] = None
                total_failed += 1
                print(f"✗ 处理错误: {e}")
        
        print("\n" + "=" * 60)
        print("批量修复完成！")
        print(f"成功: {total_success} 个文件")
        print(f"失败: {total_failed} 个文件")
        print("=" * 60)
        
        return results
