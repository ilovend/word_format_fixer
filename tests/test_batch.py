#!/usr/bin/env python3
"""
批量处理功能测试脚本
"""

import os
import tempfile
from word_format_fixer.core.fixer import RobustWordFixer


def test_batch_processing():
    """测试批量处理功能"""
    print("=" * 60)
    print("测试批量处理功能")
    print("=" * 60)
    
    # 创建临时测试文件
    test_files = []
    with tempfile.TemporaryDirectory() as temp_dir:
        # 创建3个测试文件
        for i in range(3):
            file_path = os.path.join(temp_dir, f"test_{i+1}.docx")
            # 创建一个简单的Word文档
            from docx import Document
            doc = Document()
            doc.add_heading(f"测试文档 {i+1}", 0)
            doc.add_paragraph("1. 这是一个测试段落")
            doc.add_paragraph("· 这是一个带点号的列表项")
            doc.add_table(2, 2).cell(0, 0).text = "测试表格"
            doc.save(file_path)
            test_files.append(file_path)
        
        print(f"创建了 {len(test_files)} 个测试文件:")
        for file in test_files:
            print(f"  - {file}")
        
        # 测试批量处理
        fixer = RobustWordFixer()
        output_dir = os.path.join(temp_dir, "output")
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"\n开始批量修复，输出目录: {output_dir}")
        results = fixer.fix_batch(test_files, output_dir)
        
        print(f"\n批量修复完成，结果:")
        success_count = 0
        for input_file, output_file in results.items():
            if output_file:
                success_count += 1
                print(f"  ✓ {os.path.basename(input_file)} -> {os.path.basename(output_file)}")
                # 验证输出文件存在
                assert os.path.exists(output_file), f"输出文件不存在: {output_file}"
            else:
                print(f"  ✗ {os.path.basename(input_file)} -> 失败")
        
        print(f"\n成功修复: {success_count}/{len(test_files)} 个文件")
        assert success_count == len(test_files), f"批量修复失败，成功数: {success_count}, 总数: {len(test_files)}"
        
        print("\n批量处理功能测试通过！")
        print("=" * 60)


if __name__ == "__main__":
    test_batch_processing()
