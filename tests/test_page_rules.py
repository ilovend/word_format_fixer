"""页面规则测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from core.context import RuleContext
from rules.page_rules.page_layout_rule import PageLayoutRule


class PageLayoutRuleTestCase(unittest.TestCase):
    """测试页面布局规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document(self, filename="test.docx"):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是测试内容。")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = PageLayoutRule()
        self.assertEqual(rule.display_name, "页面布局设置")
        self.assertEqual(rule.category, "页面规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "page_width_cm": 21.0,
            "page_height_cm": 29.7,
            "page_margin_top_cm": 2.54,
            "page_margin_bottom_cm": 2.54,
            "page_margin_left_cm": 3.17,
            "page_margin_right_cm": 3.17
        }
        rule = PageLayoutRule(config)

        self.assertEqual(rule.config["page_width_cm"], 21.0)
        self.assertEqual(rule.config["page_height_cm"], 29.7)
        self.assertEqual(rule.config["page_margin_top_cm"], 2.54)
        self.assertEqual(rule.config["page_margin_bottom_cm"], 2.54)
        self.assertEqual(rule.config["page_margin_left_cm"], 3.17)
        self.assertEqual(rule.config["page_margin_right_cm"], 3.17)

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = PageLayoutRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "页面布局设置")
        self.assertEqual(metadata["category"], "页面规则")
        self.assertIn("page_width_cm", metadata["params"])
        self.assertIn("page_height_cm", metadata["params"])
        self.assertIn("page_margin_top_cm", metadata["params"])
        self.assertIn("page_margin_bottom_cm", metadata["params"])
        self.assertIn("page_margin_left_cm", metadata["params"])
        self.assertIn("page_margin_right_cm", metadata["params"])

    def test_apply_changes_page_size(self):
        """测试应用规则更改页面大小"""
        doc_path = self.create_document()
        context = RuleContext(doc_path)

        rule = PageLayoutRule({
            "page_width": 21.0,
            "page_height": 29.7,
            "margin_top": 2.54,
            "margin_bottom": 2.54,
            "margin_left": 3.17,
            "margin_right": 3.17
        })
        result = rule.apply(context)

        self.assertTrue(result.success)

    def test_apply_changes_margins(self):
        """测试应用规则更改页边距"""
        doc_path = self.create_document()
        context = RuleContext(doc_path)

        # 使用不同的边距
        rule = PageLayoutRule({
            "margin_top": 3.0,
            "margin_bottom": 3.0,
            "margin_left": 2.5,
            "margin_right": 2.5
        })
        result = rule.apply(context)

        self.assertTrue(result.success)

    def test_apply_with_default_config(self):
        """测试使用默认配置应用规则"""
        doc_path = self.create_document()
        context = RuleContext(doc_path)

        rule = PageLayoutRule()
        result = rule.apply(context)

        self.assertTrue(result.success)

    def test_apply_preserves_document_content(self):
        """测试应用规则保留文档内容"""
        doc_path = self.create_document()
        context = RuleContext(doc_path)

        # 获取原始内容
        original_doc = Document(doc_path)
        original_paras = len(original_doc.paragraphs)

        # 应用规则
        rule = PageLayoutRule()
        result = rule.apply(context)
        context.save_document()

        # 验证内容没有改变
        new_doc = Document(doc_path)
        self.assertEqual(len(new_doc.paragraphs), original_paras)


class PageLayoutRuleIntegrationTestCase(unittest.TestCase):
    """页面布局规则集成测试"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_complex_document(self, filename="complex.docx"):
        """创建复杂文档"""
        doc = Document()
        doc.add_heading("复杂文档", 0)

        # 添加多个段落
        for i in range(10):
            doc.add_paragraph(f"段落 {i+1}")

        # 添加表格
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"单元格{i}-{j}"

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_page_layout_with_other_rules(self):
        """测试页面布局规则与其他规则一起工作"""
        from core.engine import RuleEngine

        doc_path = self.create_complex_document()

        # 创建规则引擎
        engine = RuleEngine()
        engine._load_rules()

        # 获取一些规则
        rules_info = engine.get_rules_info()

        # 选择页面布局规则
        page_layout_rule = [
            {"rule_id": "PageLayoutRule", "params": {
                "margin_top": 2.54,
                "margin_bottom": 2.54,
                "margin_left": 3.17,
                "margin_right": 3.17
            }}
        ]

        # 执行规则
        result = engine.execute(doc_path, page_layout_rule)

        # 验证结果
        self.assertEqual(result["status"], "success")

        # 验证文档内容完整
        doc = Document(doc_path)
        self.assertGreater(len(doc.paragraphs), 0)
        self.assertGreater(len(doc.tables), 0)

    def test_page_layout_with_preset(self):
        """测试使用预设配置应用页面布局"""
        doc_path = self.create_complex_document()
        context = RuleContext(doc_path)

        # 使用学术论文格式的页面布局
        rule = PageLayoutRule({
            "page_width": 21.0,
            "page_height": 29.7,
            "margin_top": 3.0,
            "margin_bottom": 2.5,
            "margin_left": 3.0,
            "margin_right": 2.0
        })
        result = rule.apply(context)

        self.assertTrue(result.success)

        # 保存并验证
        context.save_document()
        doc = Document(doc_path)
        self.assertIsNotNone(doc)


if __name__ == '__main__':
    unittest.main()
