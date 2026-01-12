"""性能基准测试"""

import time
from docx import Document
from word_format_fixer.core.fixer import RobustWordFixer
import os


def create_test_document(filename, paragraphs=100, tables=5):
    """
    创建测试文档
    
    Args:
        filename: 文件名
        paragraphs: 段落数
        tables: 表格数
    """
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档', level=1)
    
    # 添加段落
    for i in range(paragraphs):
        if i % 10 == 0:
            doc.add_heading(f'章节 {i//10 + 1}', level=2)
        doc.add_paragraph(f'这是第 {i+1} 个测试段落。' * 5)
    
    # 添加表格
    for i in range(tables):
        doc.add_heading(f'表格 {i+1}', level=3)
        table = doc.add_table(rows=5, cols=3)
        for r in range(5):
            for c in range(3):
                table.cell(r, c).text = f'单元格 {r+1},{c+1}'
    
    doc.save(filename)


def test_performance():
    """
    测试性能
    """
    test_file = 'test_performance.docx'
    output_file = 'test_performance_fixed.docx'
    
    # 创建测试文档
    print(f"创建测试文档...")
    create_test_document(test_file)
    
    # 初始化修复器
    fixer = RobustWordFixer()
    
    # 测试修复时间
    print(f"开始修复文档...")
    start_time = time.time()
    
    try:
        result = fixer.fix_all(test_file, output_file)
        end_time = time.time()
        elapsed = end_time - start_time
        
        print(f"修复完成！")
        print(f"输入文件大小: {os.path.getsize(test_file) / 1024:.2f} KB")
        print(f"输出文件大小: {os.path.getsize(output_file) / 1024:.2f} KB")
        print(f"处理时间: {elapsed:.2f} 秒")
        print(f"处理速度: {os.path.getsize(test_file) / elapsed / 1024:.2f} KB/秒")
        
        return elapsed
    finally:
        # 清理测试文件
        if os.path.exists(test_file):
            os.remove(test_file)
        if os.path.exists(output_file):
            os.remove(output_file)


if __name__ == '__main__':
    test_performance()
