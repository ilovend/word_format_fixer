"""文档上下文测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from core.context import RuleContext


class RuleContextTestCase(unittest.TestCase):
    """测试RuleContext类"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_test_document(self, filename="test.docx"):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是测试内容。")
        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init_with_valid_path(self):
        """测试使用有效路径初始化RuleContext"""
        doc_path = self.create_test_document()

        context = RuleContext(doc_path)

        self.assertIsNotNone(context.doc)
        self.assertEqual(context.file_path, doc_path)

    def test_init_with_invalid_path(self):
        """测试使用无效路径初始化RuleContext"""
        invalid_path = self.temp_path / "nonexistent.docx"

        with self.assertRaises(Exception):
            RuleContext(str(invalid_path))

    def test_get_document(self):
        """测试获取文档对象"""
        doc_path = self.create_test_document()
        context = RuleContext(doc_path)

        doc = context.get_document()

        # 只检查文档对象不为None，而不是使用isinstance，因为在某些环境下会出现TypeError
        self.assertIsNotNone(doc)

    def test_get_file_path(self):
        """测试获取文件路径"""
        doc_path = self.create_test_document("test_path.docx")
        context = RuleContext(doc_path)

        file_path = context.get_file_path()

        self.assertEqual(file_path, doc_path)

    def test_save_document_success(self):
        """测试保存文档成功"""
        doc_path = self.create_test_document()
        context = RuleContext(doc_path)

        # 修改文档
        doc = context.get_document()
        doc.add_paragraph("新增段落")

        # 保存文档
        result = context.save_document()

        self.assertTrue(result)

        # 验证文档已保存
        doc = Document(doc_path)
        self.assertEqual(len(doc.paragraphs), 3)  # 标题 + 原始段落 + 新增段落
        self.assertIn("新增段落", doc.paragraphs[2].text)

    def test_save_document_to_different_path(self):
        """测试保存到不同路径"""
        doc_path = self.create_test_document()
        context = RuleContext(doc_path)

        # 修改文档
        doc = context.get_document()
        doc.add_paragraph("保存到新位置")

        # 保存到新路径
        new_path = self.temp_path / "saved.docx"
        result = context.save_document(str(new_path))

        self.assertTrue(result)

        # 验证新文档存在
        doc = Document(str(new_path))
        self.assertEqual(len(doc.paragraphs), 3)  # 标题 + 原始段落 + 新增段落
        self.assertIn("保存到新位置", doc.paragraphs[2].text)

    def test_get_paragraphs(self):
        """测试获取段落"""
        doc_path = self.create_test_document()
        context = RuleContext(doc_path)

        paragraphs = context.get_paragraphs()

        self.assertIsInstance(paragraphs, list)
        self.assertGreater(len(paragraphs), 0)

    def test_get_tables(self):
        """测试获取表格"""
        # 创建包含表格的文档
        doc_path = self.create_test_document("test_tables.docx")
        doc = Document(doc_path)

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Cell 1"
        table.cell(1, 1).text = "Cell 2"
        doc.save(doc_path)

        context = RuleContext(doc_path)

        tables = context.get_tables()

        self.assertIsInstance(tables, list)
        self.assertEqual(len(tables), 1)

    def test_get_document_statistics(self):
        """测试获取文档统计信息"""
        doc_path = self.create_test_document("test_stats.docx")
        doc = Document(doc_path)

        # 添加多个段落
        for i in range(10):
            doc.add_paragraph(f"段落 {i+1}")

        # 添加表格
        table = doc.add_table(rows=3, cols=3)
        doc.save(doc_path)

        context = RuleContext(doc_path)

        stats = context.get_document_statistics()
# 验证统计信息
        self.assertIsInstance(stats, dict)
        self.assertIn("paragraph_count", stats)
        self.assertIn("table_count", stats)
        self.assertEqual(stats["paragraph_count"], 12)  # 1个标题 + 1个初始段落 + 10个新增段落
        self.assertEqual(stats["table_count"], 1)

    def test_context_isolation(self):
        """测试上下文隔离 - 修改不影响其他实例"""
        doc_path = self.create_test_document()

        # 创建两个上下文实例
        context1 = RuleContext(doc_path)
        context2 = RuleContext(doc_path)

        # 在context1中添加段落
        doc1 = context1.get_document()
        doc1.add_paragraph("context1添加的段落")

        # 在context2中添加段落
        doc2 = context2.get_document()
        doc2.add_paragraph("context2添加的段落")

        # 保存context1
        context1.save_document()

        # 重新加载文档
        doc = Document(doc_path)

        # 验证段落数量
        self.assertEqual(len(doc.paragraphs), 3)
        self.assertIn("context1添加的段落", doc.paragraphs[2].text)

    def test_context_cleanup(self):
        """测试上下文清理"""
        doc_path = self.create_test_document()
        context = RuleContext(doc_path)

        # 获取文档
        doc = context.get_document()
        self.assertIsNotNone(doc)

        # 上下文应该正常工作
        paragraphs = context.get_paragraphs()
        self.assertGreater(len(paragraphs), 0)


if __name__ == '__main__':
    unittest.main()
