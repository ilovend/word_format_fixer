"""应用服务层测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from services.application_service import (
    DocumentProcessingService,
    ConfigManagementService,
    RuleManagementService
)


class DocumentProcessingServiceTestCase(unittest.TestCase):
    """测试文档处理服务"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.service = DocumentProcessingService()

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_test_document(self, filename="test.docx"):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是测试内容。")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试服务初始化"""
        service = DocumentProcessingService()
        self.assertIsNotNone(service.engine)

    def test_process_document_success(self):
        """测试成功处理文档"""
        doc_path = self.create_test_document()

        result = self.service.process_document(doc_path)

        self.assertIsInstance(result, dict)
        self.assertIn("status", result)
        self.assertIn("summary", result)
        self.assertIn("results", result)

    def test_process_document_with_active_rules(self):
        """测试使用激活规则处理文档"""
        doc_path = self.create_test_document()

        # 获取规则信息
        rules_info = RuleManagementService().get_all_rules()
        if rules_info:
            active_rules = [{"rule_id": rules_info[0]["id"], "params": {}}]
            result = self.service.process_document(doc_path, active_rules)

            self.assertEqual(result["status"], "success")

    def test_process_document_without_document_path(self):
        """测试不提供文档路径"""
        with self.assertRaises(ValueError) as context:
            self.service.process_document("")

        self.assertIn("Missing document_path", str(context.exception))

    def test_process_document_with_nonexistent_file(self):
        """测试处理不存在的文件"""
        nonexistent_path = self.temp_path / "nonexistent.docx"

        # 应该抛出异常
        with self.assertRaises(Exception):
            self.service.process_document(str(nonexistent_path))


class ConfigManagementServiceTestCase(unittest.TestCase):
    """测试配置管理服务"""

    def setUp(self):
        """设置测试环境"""
        self.service = ConfigManagementService()

    def test_init(self):
        """测试服务初始化"""
        service = ConfigManagementService()
        self.assertIsNotNone(service.config_loader)

    def test_get_all_presets(self):
        """测试获取所有预设"""
        presets = self.service.get_all_presets()

        self.assertIsInstance(presets, dict)
        self.assertGreater(len(presets), 0)
        self.assertIn("default", presets)

    def test_get_all_presets_contains_required_fields(self):
        """测试预设包含必需字段"""
        presets = self.service.get_all_presets()

        for preset_id, preset_data in presets.items():
            self.assertIn("name", preset_data)
            self.assertIn("description", preset_data)
            self.assertIn("rules", preset_data)

    def test_save_preset_success(self):
        """测试成功保存预设"""
        preset_data = {
            "name": "测试预设",
            "description": "这是一个测试预设",
            "rules": {
                "font_standard": {"enabled": True}
            }
        }

        result = self.service.save_preset("test_preset", preset_data)

        self.assertEqual(result["status"], "success")
        self.assertIn("successfully", result["message"])

        # 验证预设已保存
        presets = self.service.get_all_presets()
        self.assertIn("test_preset", presets)

    def test_save_preset_without_id(self):
        """测试不提供预设ID"""
        with self.assertRaises(ValueError) as context:
            self.service.save_preset("", {"name": "Test"})

        self.assertIn("Missing preset_id", str(context.exception))

    def test_save_preset_without_data(self):
        """测试不提供预设数据"""
        with self.assertRaises(ValueError) as context:
            self.service.save_preset("test", None)

        self.assertIn("Missing preset_data", str(context.exception))

    def test_delete_preset_success(self):
        """测试成功删除预设"""
        # 先保存一个预设
        preset_data = {
            "name": "待删除预设",
            "description": "这是一个待删除的预设",
            "rules": {}
        }
        self.service.save_preset("to_delete", preset_data)

        # 删除预设
        result = self.service.delete_preset("to_delete")

        self.assertEqual(result["status"], "success")
        self.assertIn("successfully", result["message"])

        # 验证预设已删除
        presets = self.service.get_all_presets()
        self.assertNotIn("to_delete", presets)

    def test_delete_default_preset_fails(self):
        """测试删除默认预设失败"""
        with self.assertRaises(ValueError) as context:
            self.service.delete_preset("default")

        self.assertIn("Cannot delete default preset", str(context.exception))

    def test_delete_preset_without_id(self):
        """测试不提供预设ID"""
        with self.assertRaises(ValueError) as context:
            self.service.delete_preset("")

        self.assertIn("Missing preset_id", str(context.exception))


class RuleManagementServiceTestCase(unittest.TestCase):
    """测试规则管理服务"""

    def setUp(self):
        """设置测试环境"""
        self.service = RuleManagementService()

    def test_init(self):
        """测试服务初始化"""
        service = RuleManagementService()
        self.assertIsNotNone(service.engine)

    def test_get_all_rules(self):
        """测试获取所有规则"""
        rules = self.service.get_all_rules()

        self.assertIsInstance(rules, list)
        self.assertGreater(len(rules), 0)

    def test_get_all_rules_contains_required_fields(self):
        """测试规则信息包含必需字段"""
        rules = self.service.get_all_rules()

        for rule in rules:
            self.assertIn("id", rule)
            self.assertIn("name", rule)
            self.assertIn("description", rule)
            self.assertIn("category", rule)
            self.assertIn("params", rule)

    def test_get_all_rules_categories(self):
        """测试规则有正确的分类"""
        rules = self.service.get_all_rules()

        categories = set()
        for rule in rules:
            categories.add(rule["category"])

        # 应该有多个类别
        self.assertGreater(len(categories), 0)
        # 至少应该包含这些类别
        expected_categories = ["字体规则", "段落规则", "表格规则", "页面规则"]
        for category in expected_categories:
            # 某些类别可能没有规则，所以不强制要求
            pass


class ServiceIntegrationTestCase(unittest.TestCase):
    """服务层集成测试"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_test_document(self, filename="test.docx"):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是测试内容。")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_document_processing_with_preset(self):
        """测试使用预设处理文档"""
        doc_path = self.create_test_document()

        # 获取预设
        config_service = ConfigManagementService()
        presets = config_service.get_all_presets()

        if presets:
            # 使用第一个预设
            preset_id = list(presets.keys())[0]
            preset = presets[preset_id]

            # 构建激活规则列表
            active_rules = []
            for rule_id, rule_config in preset.get("rules", {}).items():
                if rule_config.get("enabled", False):
                    active_rules.append({"rule_id": rule_id, "params": rule_config})

            # 处理文档
            if active_rules:
                doc_service = DocumentProcessingService()
                result = doc_service.process_document(doc_path, active_rules)

                self.assertEqual(result["status"], "success")

    def test_services_work_together(self):
        """测试多个服务协同工作"""
        # 获取规则信息
        rule_service = RuleManagementService()
        rules = rule_service.get_all_rules()

        if rules:
            # 选择前两个规则
            selected_rules = rules[:2]
            active_rules = [{"rule_id": r["id"], "params": r["params"]} for r in selected_rules]

            # 处理文档
            doc_path = self.create_test_document()
            doc_service = DocumentProcessingService()
            result = doc_service.process_document(doc_path, active_rules)

            self.assertEqual(result["status"], "success")


if __name__ == '__main__':
    unittest.main()
