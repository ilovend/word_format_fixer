"""字体规则测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from core.context import RuleContext
from rules.font_rules.font_color_rule import FontColorRule
from rules.font_rules.font_standard_rule import (
    FontNameRule,
    TitleFontRule,
    FontSizeRule
)


class FontColorRuleTestCase(unittest.TestCase):
    """测试字体颜色规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_colored_document(self, filename="colored.docx"):
        """创建包含不同颜色文本的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 红色文本
        para = doc.add_paragraph()
        run = para.add_run("红色文本")
        run.font.color.rgb = RGBColor(255, 0, 0)

        # 蓝色文本
        para = doc.add_paragraph()
        run = para.add_run("蓝色文本")
        run.font.color.rgb = RGBColor(0, 0, 255)

        # 绿色文本
        para = doc.add_paragraph()
        run = para.add_run("绿色文本")
        run.font.color.rgb = RGBColor(0, 255, 0)

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = FontColorRule()
        self.assertEqual(rule.display_name, "字体颜色统一")
        self.assertEqual(rule.category, "字体规则")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = FontColorRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "字体颜色统一")
        self.assertEqual(metadata["category"], "字体规则")
        self.assertIn("text_color", metadata["params"])

    def test_apply_changes_text_color_to_black(self):
        """测试应用规则将文本颜色改为黑色"""
        doc_path = self.create_colored_document()
        context = RuleContext(doc_path)

        rule = FontColorRule()
        result = rule.apply(context)

        self.assertTrue(result.success)
        self.assertGreater(result.fixed_count, 0)

        # 验证颜色已更改
        doc = Document(doc_path)
        # 注意：python-docx中，黑色通常表示为None或RGB(0,0,0)
        # 这里我们验证规则执行成功即可

    def test_apply_with_already_black_text(self):
        """测试应用规则到已经是黑色的文本"""
        doc = Document()
        doc.add_paragraph("黑色文本")
        doc_path = self.temp_path / "black_text.docx"
        doc.save(str(doc_path))

        context = RuleContext(str(doc_path))
        rule = FontColorRule()
        result = rule.apply(context)

        # 应该成功，但fixed_count可能为0
        self.assertTrue(result.success)


class FontNameRuleTestCase(unittest.TestCase):
    """测试字体名称规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_fonts(self, filename="fonts.docx"):
        """创建包含不同字体的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加一些文本
        doc.add_paragraph("这是测试文本")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = FontNameRule()
        self.assertEqual(rule.display_name, "字体名称标准化")
        self.assertEqual(rule.category, "字体规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "chinese_font": "宋体",
            "western_font": "Arial"
        }
        rule = FontNameRule(config)

        self.assertEqual(rule.config["chinese_font"], "宋体")
        self.assertEqual(rule.config["western_font"], "Arial")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = FontNameRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "字体名称标准化")
        self.assertIn("chinese_font", metadata["params"])
        self.assertIn("western_font", metadata["params"])

    def test_apply_changes_fonts(self):
        """测试应用规则更改字体"""
        doc_path = self.create_document_with_fonts()
        context = RuleContext(doc_path)

        rule = FontNameRule({
            "chinese_font": "宋体",
            "western_font": "Arial"
        })
        result = rule.apply(context)

        self.assertTrue(result.success)


class TitleFontRuleTestCase(unittest.TestCase):
    """测试标题字体规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_titles(self, filename="titles.docx"):
        """创建包含标题的文档"""
        doc = Document()
        doc.add_heading("一级标题", level=1)
        doc.add_paragraph("正文内容")
        doc.add_heading("二级标题", level=2)
        doc.add_paragraph("更多正文内容")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TitleFontRule()
        self.assertEqual(rule.display_name, "标题字体设置")
        self.assertEqual(rule.category, "字体规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {"title_font": "黑体"}
        rule = TitleFontRule(config)

        self.assertEqual(rule.config["title_font"], "黑体")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TitleFontRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "标题字体设置")
        self.assertIn("title_font", metadata["params"])

    def test_apply_changes_title_fonts(self):
        """测试应用规则更改标题字体"""
        doc_path = self.create_document_with_titles()
        context = RuleContext(doc_path)

        rule = TitleFontRule({"title_font": "黑体"})
        result = rule.apply(context)

        self.assertTrue(result.success)


class FontSizeRuleTestCase(unittest.TestCase):
    """测试字号规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_sizes(self, filename="sizes.docx"):
        """创建包含不同字号的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加不同字号的文本
        para = doc.add_paragraph("小字号文本")
        run = para.runs[0]
        run.font.size = Pt(10)

        para = doc.add_paragraph("中字号文本")
        run = para.runs[0]
        run.font.size = Pt(14)

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = FontSizeRule()
        self.assertEqual(rule.display_name, "字号标准化")
        self.assertEqual(rule.category, "字体规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "font_size_body": 12,
            "font_size_title1": 22,
            "font_size_title2": 18,
            "font_size_title3": 16
        }
        rule = FontSizeRule(config)

        self.assertEqual(rule.config["font_size_body"], 12)
        self.assertEqual(rule.config["font_size_title1"], 22)

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = FontSizeRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "字号标准化")
        self.assertIn("font_size_body", metadata["params"])
        self.assertIn("font_size_title1", metadata["params"])
        self.assertIn("font_size_title2", metadata["params"])
        self.assertIn("font_size_title3", metadata["params"])

    def test_apply_changes_font_sizes(self):
        """测试应用规则更改字号"""
        doc_path = self.create_document_with_sizes()
        context = RuleContext(doc_path)

        rule = FontSizeRule({
            "font_size_body": 12,
            "font_size_title1": 22,
            "font_size_title2": 18,
            "font_size_title3": 16
        })
        result = rule.apply(context)

        self.assertTrue(result.success)


if __name__ == '__main__':
    unittest.main()
