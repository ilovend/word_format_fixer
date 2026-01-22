"""测试横线移除规则"""

import unittest
import os
from unittest.mock import MagicMock, patch
import sys
import tempfile
from docx import Document

# 获取项目根目录
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# 添加python-backend目录到Python路径
sys.path.insert(0, os.path.join(project_root, 'python-backend'))

# 导入需要测试的规则
from rules.paragraph_rules.horizontal_rule_removal_rule import HorizontalRuleRemovalRule
from core.context import RuleContext
from rules.base_rule import RuleResult


class TestHorizontalRuleRemovalRule(unittest.TestCase):
    """测试横线移除规则"""

    def setUp(self):
        """设置测试环境"""
        self.rule = HorizontalRuleRemovalRule()
    
    def test_init(self):
        """测试规则初始化"""
        self.assertIsInstance(self.rule, HorizontalRuleRemovalRule)
        self.assertEqual(self.rule.display_name, "横线移除")
        self.assertEqual(self.rule.category, "段落规则")
    
    def test_get_metadata(self):
        """测试获取规则元数据"""
        metadata = self.rule.get_metadata()
        self.assertEqual(metadata["name"], "横线移除")
        self.assertEqual(metadata["category"], "段落规则")
        self.assertIn("remove_horizontal_rules", metadata["params"])
    
    def create_test_document(self, temp_dir):
        """创建包含横线的测试文档"""
        # 创建一个新文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('测试文档', 0)
        
        # 添加一些文本
        doc.add_paragraph('这是测试文档的第一段内容。')
        
        # 添加不同类型的横线
        doc.add_paragraph('---')  # 连字符横线
        doc.add_paragraph('这是横线后的内容。')
        
        doc.add_paragraph('***')  # 星号横线
        doc.add_paragraph('这是另一条横线后的内容。')
        
        doc.add_paragraph('___')  # 下划线横线
        doc.add_paragraph('这是第三条横线后的内容。')
        
        doc.add_paragraph('- - -')  # 带空格的连字符横线
        doc.add_paragraph('这是第四条横线后的内容。')
        
        doc.add_paragraph('* * *')  # 带空格的星号横线
        doc.add_paragraph('这是第五条横线后的内容。')
        
        doc.add_paragraph('_ _ _')  # 带空格的下划线横线
        doc.add_paragraph('这是第六条横线后的内容。')
        
        # 保存文档
        test_doc_path = os.path.join(temp_dir, 'test_horizontal_rules.docx')
        doc.save(test_doc_path)
        return test_doc_path
    
    def test_apply(self):
        """测试规则应用"""
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 创建测试文档
            test_doc_path = self.create_test_document(temp_dir)
            
            # 创建文档上下文
            context = RuleContext(test_doc_path)
            
            # 执行规则
            result = self.rule.apply(context)
            
            # 验证结果
            self.assertIsInstance(result, RuleResult)
            self.assertTrue(result.success)
            self.assertEqual(result.fixed_count, 6)  # 应该移除6个横线
            self.assertIn("移除了6个横线段落", result.details)
            
            # 保存并重新加载文档，验证横线确实被移除
            context.save_document()
            doc = Document(test_doc_path)
            
            # 检查文档中是否还有横线
            horizontal_rule_patterns = [
                '---', '***', '___', '- - -', '* * *', '_ _ _'
            ]
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                for pattern in horizontal_rule_patterns:
                    self.assertNotEqual(text, pattern)
    
    def test_no_horizontal_rules(self):
        """测试文档中没有横线的情况"""
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 创建不包含横线的测试文档
            test_doc_path = os.path.join(temp_dir, 'test_no_horizontal_rules.docx')
            doc = Document()
            doc.add_heading('测试文档', 0)
            doc.add_paragraph('这是测试文档的内容，没有横线。')
            doc.add_paragraph('这是另一段内容。')
            doc.save(test_doc_path)
            
            # 创建文档上下文
            context = RuleContext(test_doc_path)
            
            # 执行规则
            result = self.rule.apply(context)
            
            # 验证结果
            self.assertIsInstance(result, RuleResult)
            self.assertTrue(result.success)
            self.assertEqual(result.fixed_count, 0)  # 应该移除0个横线
            self.assertIn("文档中没有需要移除的横线", result.details)
    
    def test_rule_in_rule_engine(self):
        """测试规则引擎能否正确加载和执行横线移除规则"""
        from core.engine import RuleEngine
        
        engine = RuleEngine()
        
        # 检查规则是否被正确加载
        rule_id = "HorizontalRuleRemovalRule"
        self.assertIn(rule_id, engine.rules)
        
        # 创建临时目录和测试文档
        with tempfile.TemporaryDirectory() as temp_dir:
            test_doc_path = self.create_test_document(temp_dir)
            
            # 执行规则
            result = engine.execute(test_doc_path, [{'rule_id': rule_id, 'params': {}}])
            
            # 验证结果
            self.assertEqual(result['status'], 'success')
            self.assertEqual(result['summary']['total_fixed'], 6)
            self.assertEqual(len(result['results']), 1)
            self.assertEqual(result['results'][0]['rule_id'], rule_id)
            self.assertTrue(result['results'][0]['success'])
            self.assertEqual(result['results'][0]['fixed_count'], 6)


if __name__ == '__main__':
    unittest.main()
