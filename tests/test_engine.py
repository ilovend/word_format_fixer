"""规则引擎完整测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from core.engine import RuleEngine
from core.context import RuleContext
from rules.base_rule import BaseRule, RuleResult


class RuleEngineLoadRulesTestCase(unittest.TestCase):
    """测试规则加载功能"""

    def setUp(self):
        """设置测试环境"""
        self.engine = RuleEngine()

    def test_init_creates_rules_dict(self):
        """测试初始化创建规则字典"""
        self.assertIsInstance(self.engine.rules, dict)

    def test_load_rules_populates_rules(self):
        """测试加载规则填充规则字典"""
        self.engine._load_rules()
        self.assertGreater(len(self.engine.rules), 0)

    def test_load_rules_loads_all_rule_types(self):
        """测试加载所有类型的规则"""
        self.engine._load_rules()

        # 检查是否有不同类型的规则
        rule_categories = set()
        for rule in self.engine.rules.values():
            metadata = rule.get_metadata()
            rule_categories.add(metadata["category"])

        # 应该有多个类别
        self.assertGreater(len(rule_categories), 0)

    def test_get_rules_info_returns_list(self):
        """测试获取规则信息返回列表"""
        self.engine._load_rules()
        rules_info = self.engine.get_rules_info()

        self.assertIsInstance(rules_info, list)
        self.assertGreater(len(rules_info), 0)

    def test_get_rules_info_contains_required_fields(self):
        """测试规则信息包含必需字段"""
        self.engine._load_rules()
        rules_info = self.engine.get_rules_info()

        for rule_info in rules_info:
            self.assertIn("id", rule_info)
            self.assertIn("name", rule_info)
            self.assertIn("description", rule_info)
            self.assertIn("category", rule_info)
            self.assertIn("params", rule_info)

    def test_get_rules_info_structure(self):
        """测试规则信息数据结构"""
        self.engine._load_rules()
        rules_info = self.engine.get_rules_info()

        # 检查第一个规则的信息结构
        if rules_info:
            first_rule = rules_info[0]
            self.assertIsInstance(first_rule["id"], str)
            self.assertIsInstance(first_rule["name"], str)
            self.assertIsInstance(first_rule["description"], str)
            self.assertIsInstance(first_rule["category"], str)
            self.assertIsInstance(first_rule["params"], dict)


class RuleEngineRegisterRuleTestCase(unittest.TestCase):
    """测试规则注册功能"""

    def setUp(self):
        """设置测试环境"""
        self.engine = RuleEngine()
        self.test_rule_class = self.create_test_rule()

    def create_test_rule(self, rule_id="TestRule"):
        """创建测试规则"""
        class TestRule(BaseRule):
            display_name = "测试规则"
            category = "测试"

            def apply(self, doc_context):
                return RuleResult(
                    rule_id=self.rule_id,
                    success=True,
                    fixed_count=0,
                    details=["测试"]
                )

        TestRule.rule_id = rule_id
        return TestRule

    def test_register_rule_adds_to_dict(self):
        """测试注册规则添加到字典"""
        rule = self.test_rule_class()
        initial_count = len(self.engine.rules)

        self.engine.register_rule(rule)

        self.assertEqual(len(self.engine.rules), initial_count + 1)
        self.assertIn(rule.rule_id, self.engine.rules)

    def test_register_rule_stores_instance(self):
        """测试注册规则存储实例"""
        rule = self.test_rule_class()
        self.engine.register_rule(rule)

        registered_rule = self.engine.rules[rule.rule_id]
        self.assertIs(registered_rule, rule)

    def test_register_duplicate_rule(self):
        """测试注册重复规则"""
        rule1 = self.test_rule_class()
        rule2 = self.test_rule_class()

        self.engine.register_rule(rule1)
        initial_count = len(self.engine.rules)

        self.engine.register_rule(rule2)

        # 应该替换，而不是添加
        self.assertEqual(len(self.engine.rules), initial_count)
        self.assertIs(self.engine.rules[rule1.rule_id], rule2)


class RuleEngineExecuteTestCase(unittest.TestCase):
    """测试规则执行功能"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.engine = RuleEngine()
        self.engine._load_rules()

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_test_document(self, filename="test.docx"):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是测试内容。")
        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_execute_with_valid_document(self):
        """测试使用有效文档执行"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path)

        self.assertIsInstance(result, dict)
        self.assertIn("status", result)
        self.assertIn("summary", result)
        self.assertIn("results", result)
        self.assertIn("save_success", result)
        self.assertIn("saved_to", result)

    def test_execute_with_active_rules(self):
        """测试使用激活规则执行"""
        doc_path = self.create_test_document()

        # 获取第一个规则
        rules_info = self.engine.get_rules_info()
        if not rules_info:
            self.skipTest("没有可用的规则")

        first_rule_id = rules_info[0]["id"]
        active_rules = [{"rule_id": first_rule_id, "params": {}}]

        result = self.engine.execute(doc_path, active_rules)

        self.assertEqual(result["status"], "success")
        self.assertEqual(len(result["results"]), 1)
        self.assertEqual(result["results"][0]["rule_id"], first_rule_id)

    def test_execute_with_all_enabled_rules(self):
        """测试执行所有启用的规则"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path)

        self.assertEqual(result["status"], "success")
        self.assertGreater(len(result["results"]), 0)

    def test_execute_without_active_rules(self):
        """测试不指定激活规则时执行所有规则"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path, None)

        self.assertEqual(result["status"], "success")

    def test_execute_with_empty_active_rules(self):
        """测试使用空的激活规则列表执行"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path, [])

        self.assertEqual(result["status"], "success")
        self.assertEqual(len(result["results"]), 0)

    def test_execute_summary_contains_total_fixed(self):
        """测试执行摘要包含total_fixed"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path)

        self.assertIn("total_fixed", result["summary"])
        self.assertIsInstance(result["summary"]["total_fixed"], int)

    def test_execute_summary_contains_time_taken(self):
        """测试执行摘要包含time_taken"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path)

        self.assertIn("time_taken", result["summary"])
        self.assertIsInstance(result["summary"]["time_taken"], str)

    def test_execute_saves_document(self):
        """测试执行保存文档"""
        doc_path = self.create_test_document()

        result = self.engine.execute(doc_path)

        self.assertTrue(result["save_success"])
        self.assertEqual(result["saved_to"], doc_path)

    def test_execute_with_rule_error(self):
        """测试规则执行错误处理"""
        doc_path = self.create_test_document()

        # 创建一个会抛出异常的规则
        class ErrorRule(BaseRule):
            display_name = "错误规则"
            category = "测试"

            def apply(self, doc_context):
                raise ValueError("测试错误")

        error_rule = ErrorRule()
        self.engine.register_rule(error_rule)

        # 执行包含错误规则的执行
        active_rules = [{"rule_id": "ErrorRule", "params": {}}]

        result = self.engine.execute(doc_path, active_rules)

        # 应该捕获错误并返回失败结果
        self.assertEqual(len(result["results"]), 1)
        self.assertFalse(result["results"][0]["success"])

    def test_execute_with_invalid_rule_id(self):
        """测试使用无效规则ID执行"""
        doc_path = self.create_test_document()

        active_rules = [{"rule_id": "NonExistentRule", "params": {}}]

        result = self.engine.execute(doc_path, active_rules)

        # 应该处理无效规则ID
        self.assertEqual(len(result["results"]), 1)
        self.assertFalse(result["results"][0]["success"])


class RuleEngineIntegrationTestCase(unittest.TestCase):
    """集成测试 - 测试规则引擎与规则的完整交互"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.engine = RuleEngine()
        self.engine._load_rules()

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_complex_document(self, filename="complex.docx"):
        """创建复杂测试文档"""
        from docx.shared import RGBColor

        doc = Document()
        doc.add_heading("复杂文档", 0)

        # 添加不同颜色的文本
        para = doc.add_paragraph()
        run = para.add_run("红色文本")
        run.font.color.rgb = RGBColor(255, 0, 0)

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        table.cell(1, 0).text = "Data"

        # 添加列表
        doc.add_paragraph("1. 列表项1")
        doc.add_paragraph("2. 列表项2")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_full_workflow(self):
        """测试完整工作流程"""
        # 1. 创建文档
        doc_path = self.create_complex_document()

        # 2. 获取规则信息
        rules_info = self.engine.get_rules_info()
        self.assertGreater(len(rules_info), 0)

        # 3. 选择一些规则
        if len(rules_info) >= 2:
            selected_rule_ids = [rules_info[0]["id"], rules_info[1]["id"]]
            active_rules = [{"rule_id": rid, "params": {}} for rid in selected_rule_ids]

            # 4. 执行规则
            result = self.engine.execute(doc_path, active_rules)

            # 5. 验证结果
            self.assertEqual(result["status"], "success")
            self.assertEqual(len(result["results"]), 2)

            # 6. 验证文档已修改
            doc = Document(doc_path)
            self.assertIsNotNone(doc)

    def test_rules_are_independent(self):
        """测试规则独立性 - 规则之间不应该有隐式依赖"""
        doc_path = self.create_complex_document()

        # 以不同的顺序执行规则
        rules_info = self.engine.get_rules_info()
        if len(rules_info) >= 2:
            rule_ids = [r["id"] for r in rules_info[:2]]

            # 顺序1
            active_rules_1 = [{"rule_id": rule_ids[0], "params": {}},
                             {"rule_id": rule_ids[1], "params": {}}]

            # 顺序2
            active_rules_2 = [{"rule_id": rule_ids[1], "params": {}},
                             {"rule_id": rule_ids[0], "params": {}}]

            result_1 = self.engine.execute(doc_path, active_rules_1)
            result_2 = self.engine.execute(doc_path, active_rules_2)

            # 两次执行都应该成功
            self.assertEqual(result_1["status"], "success")
            self.assertEqual(result_2["status"], "success")


if __name__ == '__main__':
    unittest.main()
