"""
测试配置文件 - pytest配置
"""

import sys
import os
import tempfile
from pathlib import Path
import pytest

# 获取项目根目录
project_root = Path(__file__).parent.parent

# 添加python-backend目录到Python路径
python_backend_path = project_root / "python-backend"
sys.path.insert(0, str(python_backend_path))

# 设置测试环境变量
os.environ["TESTING"] = "1"


def pytest_configure(config):
    """配置pytest"""
    # 注册自定义标记
    config.addinivalue_line(
        "markers", "unit: 单元测试"
    )
    config.addinivalue_line(
        "markers", "integration: 集成测试"
    )
    config.addinivalue_line(
        "markers", "slow: 慢速测试"
    )


@pytest.fixture(scope="session")
def project_root_dir():
    """项目根目录"""
    return project_root


@pytest.fixture(scope="session")
def python_backend_dir():
    """Python后端目录"""
    return python_backend_path


@pytest.fixture
def temp_dir():
    """临时目录fixture"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_doc_path(temp_dir):
    """创建示例Word文档用于测试"""
    from docx import Document

    doc = Document()
    doc.add_heading("测试文档", 0)

    # 添加一些文本
    doc.add_paragraph("这是测试文档的第一段内容。")
    doc.add_paragraph("这是第二段内容。")

    # 添加一个标题
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("标题下的内容。")

    # 添加列表
    doc.add_paragraph("列表项1")
    doc.add_paragraph("列表项2")

    # 保存文档
    doc_path = temp_dir / "test_document.docx"
    doc.save(str(doc_path))

    return doc_path


@pytest.fixture
def sample_table_doc_path(temp_dir):
    """创建包含表格的示例Word文档用于测试"""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("测试表格文档", 0)

    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # 填充表格
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"单元格 {i+1}-{j+1}"

    # 添加一些文本
    doc.add_paragraph("表格前的文本。")
    doc.add_paragraph("表格后的文本。")

    # 保存文档
    doc_path = temp_dir / "test_table_document.docx"
    doc.save(str(doc_path))

    return doc_path


@pytest.fixture
def sample_complex_doc_path(temp_dir):
    """创建复杂的示例Word文档用于测试"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("复杂测试文档", 0)

    # 添加多种标题
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("一级标题下的内容。")

    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("二级标题下的内容。")

    doc.add_heading("三级标题", level=3)
    doc.add_paragraph("三级标题下的内容。")

    # 添加不同格式的文本
    para = doc.add_paragraph()
    run = para.add_run("这是加粗文本")
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("这是红色文本")
    run.font.color.rgb = RGBColor(255, 0, 0)

    # 添加列表
    doc.add_paragraph("1. 编号列表项1")
    doc.add_paragraph("2. 编号列表项2")
    doc.add_paragraph("· 项目符号项1")
    doc.add_paragraph("· 项目符号项2")

    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头1"
    table.cell(0, 1).text = "表头2"
    table.cell(1, 0).text = "数据1"
    table.cell(1, 1).text = "数据2"

    # 保存文档
    doc_path = temp_dir / "test_complex_document.docx"
    doc.save(str(doc_path))

    return doc_path


# 导入pytest（确保模块存在）
try:
    import pytest
except ImportError:
    # 如果pytest未安装，定义一个假的pytest装饰器
    class pytest:
        @staticmethod
        def fixture(func=None, *args, **kwargs):
            def decorator(f):
                return f
            return decorator(func) if func else decorator
