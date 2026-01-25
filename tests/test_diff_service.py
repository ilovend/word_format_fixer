"""
DiffService 单元测试
测试文档对比功能
"""

import unittest
import tempfile
import os
from pathlib import Path
from docx import Document

import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'python-backend'))

from services.diff_service import DiffService


class DiffServiceTestCase(unittest.TestCase):
    """DiffService 测试用例"""

    def setUp(self):
        """测试前准备"""
        self.service = DiffService()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """测试后清理"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_test_document(self, content="测试内容"):
        """创建测试文档"""
        doc_path = os.path.join(self.temp_dir, "test.docx")
        doc = Document()
        doc.add_heading("测试标题", 0)
        doc.add_paragraph(content)
        doc.save(doc_path)
        return doc_path

    def test_get_document_preview(self):
        """测试获取文档HTML预览"""
        doc_path = self.create_test_document("这是测试段落")
        
        result = self.service.get_document_preview(doc_path)
        
        self.assertEqual(result["status"], "success")
        self.assertIn("html", result)
        self.assertIn("测试标题", result["html"])
        self.assertIn("测试段落", result["html"])

    def test_get_document_preview_nonexistent_file(self):
        """测试预览不存在的文件"""
        result = self.service.get_document_preview("/nonexistent/path.docx")
        
        self.assertEqual(result["status"], "error")
        self.assertIn("message", result)

    def test_prepare_diff(self):
        """测试准备对比"""
        doc_path = self.create_test_document()
        
        result = self.service.prepare_diff(doc_path)
        
        self.assertEqual(result["status"], "success")
        self.assertIn("original_html", result)

    def test_generate_diff_without_prepare(self):
        """测试未准备就生成对比"""
        doc_path = self.create_test_document()
        
        result = self.service.generate_diff(doc_path)
        
        self.assertEqual(result["status"], "error")
        self.assertIn("No original document cached", result["message"])

    def test_full_diff_workflow(self):
        """测试完整的对比工作流"""
        doc_path = self.create_test_document("原始内容")
        
        # 1. 准备对比
        prepare_result = self.service.prepare_diff(doc_path)
        self.assertEqual(prepare_result["status"], "success")
        
        # 2. 修改文档
        doc = Document(doc_path)
        doc.add_paragraph("新增内容")
        doc.save(doc_path)
        
        # 3. 生成对比
        diff_result = self.service.generate_diff(doc_path)
        
        self.assertEqual(diff_result["status"], "success")
        self.assertIn("original_html", diff_result)
        self.assertIn("modified_html", diff_result)
        self.assertIn("diff_html", diff_result)
        self.assertIn("stats", diff_result)
        
        # 验证统计信息
        stats = diff_result["stats"]
        self.assertIn("additions", stats)
        self.assertIn("deletions", stats)
        self.assertIn("total_changes", stats)

    def test_diff_detects_changes(self):
        """测试对比能检测到变更"""
        doc_path = self.create_test_document("原始内容ABC")
        
        # 准备对比
        self.service.prepare_diff(doc_path)
        
        # 修改文档内容
        doc = Document(doc_path)
        doc.add_paragraph("这是新增的段落XYZ")
        doc.save(doc_path)
        
        # 生成对比
        diff_result = self.service.generate_diff(doc_path)
        
        self.assertEqual(diff_result["status"], "success")
        # 修改后的HTML应该包含新内容
        self.assertIn("XYZ", diff_result["modified_html"])


if __name__ == '__main__':
    unittest.main()
