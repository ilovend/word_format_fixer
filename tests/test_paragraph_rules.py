"""段落规则测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.context import RuleContext
from rules.paragraph_rules.paragraph_spacing_rule import ParagraphSpacingRule
from rules.paragraph_rules.title_bold_rule import TitleBoldRule
from rules.paragraph_rules.title_alignment_rule import TitleAlignmentRule
from rules.paragraph_rules.list_numbering_rule import ListNumberingRule
from rules.paragraph_rules.horizontal_rule_removal_rule import HorizontalRuleRemovalRule


class ParagraphSpacingRuleTestCase(unittest.TestCase):
    """测试段落间距规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_spacing(self, filename="spacing.docx"):
        """创建包含不同段落间距的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        para = doc.add_paragraph("第一段")
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        para = doc.add_paragraph("第二段")
        para.paragraph_format.line_spacing = 2.0

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = ParagraphSpacingRule()
        self.assertEqual(rule.display_name, "段落间距统一")
        self.assertEqual(rule.category, "段落规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6
        }
        rule = ParagraphSpacingRule(config)

        self.assertEqual(rule.config["line_spacing"], 1.5)
        self.assertEqual(rule.config["space_before"], 0)
        self.assertEqual(rule.config["space_after"], 6)

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = ParagraphSpacingRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "段落间距统一")
        self.assertIn("body_line_spacing", metadata["params"])
        self.assertIn("body_space_before", metadata["params"])
        self.assertIn("body_space_after", metadata["params"])

    def test_apply_changes_spacing(self):
        """测试应用规则更改段落间距"""
        doc_path = self.create_document_with_spacing()
        context = RuleContext(doc_path)

        rule = ParagraphSpacingRule({
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6
        })
        result = rule.apply(context)

        self.assertTrue(result.success)


class TitleBoldRuleTestCase(unittest.TestCase):
    """测试标题加粗规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_titles(self, filename="titles.docx"):
        """创建包含标题的文档"""
        doc = Document()
        doc.add_heading("一级标题", level=1)
        doc.add_paragraph("正文内容")
        doc.add_heading("二级标题", level=2)
        doc.add_paragraph("更多正文内容")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TitleBoldRule()
        self.assertEqual(rule.display_name, "标题加粗")
        self.assertEqual(rule.category, "段落规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {"bold": True}
        rule = TitleBoldRule(config)

        self.assertEqual(rule.config["bold"], True)

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TitleBoldRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "标题加粗")
        self.assertIn("bold", metadata["params"])

    def test_apply_makes_titles_bold(self):
        """测试应用规则使标题加粗"""
        doc_path = self.create_document_with_titles()
        context = RuleContext(doc_path)

        rule = TitleBoldRule({"enabled": True})
        result = rule.apply(context)

        self.assertTrue(result.success)


class TitleAlignmentRuleTestCase(unittest.TestCase):
    """测试标题对齐规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_titles(self, filename="titles.docx"):
        """创建包含标题的文档"""
        doc = Document()
        doc.add_heading("一级标题", level=1)
        doc.add_paragraph("正文内容")
        doc.add_heading("二级标题", level=2)
        doc.add_paragraph("更多正文内容")
        doc.add_heading("三级标题", level=3)
        doc.add_paragraph("更多内容")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TitleAlignmentRule()
        self.assertEqual(rule.display_name, "标题对齐设置")
        self.assertEqual(rule.category, "段落规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "heading1_align": "center",
            "other_heading_align": "left"
        }
        rule = TitleAlignmentRule(config)

        self.assertEqual(rule.config["heading1_align"], "center")
        self.assertEqual(rule.config["other_heading_align"], "left")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TitleAlignmentRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "标题对齐设置")
        self.assertIn("heading1_align", metadata["params"])
        self.assertIn("other_heading_align", metadata["params"])

    def test_apply_changes_title_alignment(self):
        """测试应用规则更改标题对齐"""
        doc_path = self.create_document_with_titles()
        context = RuleContext(doc_path)

        rule = TitleAlignmentRule({
            "title1_align": "center",
            "other_titles_align": "left"
        })
        result = rule.apply(context)

        self.assertTrue(result.success)


class ListNumberingRuleTestCase(unittest.TestCase):
    """测试列表编号规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_lists(self, filename="lists.docx"):
        """创建包含列表的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加编号列表
        doc.add_paragraph("1. 列表项1")
        doc.add_paragraph("2. 列表项2")
        doc.add_paragraph("3. 列表项3")

        # 添加项目符号列表
        doc.add_paragraph("· 项目符号1")
        doc.add_paragraph("· 项目符号2")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = ListNumberingRule()
        self.assertEqual(rule.display_name, "编号列表标准化")
        self.assertEqual(rule.category, "段落规则")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = ListNumberingRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "编号列表标准化")

    def test_apply_fixes_list_numbering(self):
        """测试应用规则修复列表编号"""
        doc_path = self.create_document_with_lists()
        context = RuleContext(doc_path)

        rule = ListNumberingRule()
        result = rule.apply(context)

        self.assertTrue(result.success)


class HorizontalRuleRemovalRuleTestCase(unittest.TestCase):
    """测试横线移除规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_horizontal_rules(self, filename="hr.docx"):
        """创建包含横线的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        doc.add_paragraph("横线前文本")
        doc.add_paragraph("---")
        doc.add_paragraph("横线后文本")

        doc.add_paragraph("第二条横线前")
        doc.add_paragraph("***")
        doc.add_paragraph("第二条横线后")

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = HorizontalRuleRemovalRule()
        self.assertEqual(rule.display_name, "横线移除")
        self.assertEqual(rule.category, "段落规则")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = HorizontalRuleRemovalRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "横线移除")
        self.assertIn("remove_horizontal_rules", metadata["params"])

    def test_apply_removes_horizontal_rules(self):
        """测试应用规则移除横线"""
        doc_path = self.create_document_with_horizontal_rules()
        context = RuleContext(doc_path)

        rule = HorizontalRuleRemovalRule()
        result = rule.apply(context)

        self.assertTrue(result.success)
        self.assertGreater(result.fixed_count, 0)

    def test_apply_with_no_horizontal_rules(self):
        """测试应用到没有横线的文档"""
        doc = Document()
        doc.add_paragraph("只有文本")
        doc_path = self.temp_path / "no_hr.docx"
        doc.save(str(doc_path))

        context = RuleContext(str(doc_path))
        rule = HorizontalRuleRemovalRule()
        result = rule.apply(context)

        self.assertTrue(result.success)
        # 可能没有移除任何东西


if __name__ == '__main__':
    unittest.main()
