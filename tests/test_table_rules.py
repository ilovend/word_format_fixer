"""表格规则测试"""

import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
from core.context import RuleContext
from rules.table_rules.table_width_rule import TableWidthRule
from rules.table_rules.table_border_rule import TableBorderRule
from rules.table_rules.table_borders_rule import TableBordersRule


class TableWidthRuleTestCase(unittest.TestCase):
    """测试表格宽度规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_tables(self, filename="tables.docx"):
        """创建包含表格的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加表格1
        table1 = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table1.cell(i, j).text = f"表1-{i}-{j}"

        # 添加一些文本
        doc.add_paragraph("表格1和表格2之间的文本")

        # 添加表格2
        table2 = doc.add_table(rows=2, cols=2)
        for i in range(2):
            for j in range(2):
                table2.cell(i, j).text = f"表2-{i}-{j}"

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TableWidthRule()
        self.assertEqual(rule.display_name, "表格宽度优化")
        self.assertEqual(rule.category, "表格规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "table_width_percent": 95,
            "auto_adjust_columns": True
        }
        rule = TableWidthRule(config)

        self.assertEqual(rule.config["table_width_percent"], 95)
        self.assertEqual(rule.config["auto_adjust_columns"], True)

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TableWidthRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "表格宽度优化")
        self.assertIn("table_width_percent", metadata["params"])
        self.assertIn("auto_adjust_columns", metadata["params"])

    def test_apply_changes_table_width(self):
        """测试应用规则更改表格宽度"""
        doc_path = self.create_document_with_tables()
        context = RuleContext(doc_path)

        rule = TableWidthRule({
            "table_width_percent": 95,
            "auto_adjust_columns": True
        })
        result = rule.apply(context)

        self.assertTrue(result.success)

    def test_apply_with_no_tables(self):
        """测试应用到没有表格的文档"""
        doc = Document()
        doc.add_paragraph("没有表格的文档")
        doc_path = self.temp_path / "no_tables.docx"
        doc.save(str(doc_path))

        context = RuleContext(str(doc_path))
        rule = TableWidthRule()
        result = rule.apply(context)

        self.assertTrue(result.success)


class TableBorderRuleTestCase(unittest.TestCase):
    """测试表格边框规则"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_tables(self, filename="tables.docx"):
        """创建包含表格的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加表格
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"单元格{i}-{j}"

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TableBorderRule()
        self.assertEqual(rule.display_name, "表格边框添加")
        self.assertEqual(rule.category, "表格规则")

    def test_init_with_config(self):
        """测试使用配置初始化规则"""
        config = {
            "border_size": 1,
            "border_color": "000000"
        }
        rule = TableBorderRule(config)

        self.assertEqual(rule.config["border_size"], 1)
        self.assertEqual(rule.config["border_color"], "000000")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TableBorderRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "表格边框添加")
        self.assertIn("border_size", metadata["params"])
        self.assertIn("border_color", metadata["params"])

    def test_apply_adds_borders(self):
        """测试应用规则添加边框"""
        doc_path = self.create_document_with_tables()
        context = RuleContext(doc_path)

        rule = TableBorderRule({
            "border_size": 1,
            "border_color": "000000"
        })
        result = rule.apply(context)

        self.assertTrue(result.success)


class TableBordersRuleTestCase(unittest.TestCase):
    """测试表格边框规则（复数形式）"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_document_with_tables(self, filename="tables.docx"):
        """创建包含表格的文档"""
        doc = Document()
        doc.add_heading("测试文档", 0)

        # 添加多个表格
        for i in range(2):
            table = doc.add_table(rows=2, cols=2)
            for row in range(2):
                for col in range(2):
                    table.cell(row, col).text = f"表{i+1}-{row}-{col}"

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init(self):
        """测试规则初始化"""
        rule = TableBordersRule()
        self.assertEqual(rule.display_name, "表格边框添加")
        self.assertEqual(rule.category, "表格规则")

    def test_get_metadata(self):
        """测试获取规则元数据"""
        rule = TableBordersRule()
        metadata = rule.get_metadata()

        self.assertEqual(metadata["name"], "表格边框添加")

    def test_apply_adds_borders_to_all_tables(self):
        """测试应用规则为所有表格添加边框"""
        doc_path = self.create_document_with_tables()
        context = RuleContext(doc_path)

        rule = TableBordersRule()
        result = rule.apply(context)

        self.assertTrue(result.success)

    def test_apply_with_no_tables(self):
        """测试应用到没有表格的文档"""
        doc = Document()
        doc.add_paragraph("没有表格的文档")
        doc_path = self.temp_path / "no_tables.docx"
        doc.save(str(doc_path))

        context = RuleContext(str(doc_path))
        rule = TableBordersRule()
        result = rule.apply(context)

        self.assertTrue(result.success)


class TableRuleIntegrationTestCase(unittest.TestCase):
    """表格规则集成测试"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        """清理测试环境"""
        self.temp_dir.cleanup()

    def create_complex_table_document(self, filename="complex_tables.docx"):
        """创建包含复杂表格的文档"""
        from docx.shared import RGBColor

        doc = Document()
        doc.add_heading("复杂表格文档", 0)

        # 创建一个较大且复杂的表格
        table = doc.add_table(rows=5, cols=4)

        # 填充表格内容
        for i in range(5):
            for j in range(4):
                cell = table.cell(i, j)
                cell.text = f"单元格{i+1}-{j+1}"

        # 添加一些文本
        doc.add_paragraph("表格前的文本")
        doc.add_paragraph("表格后的文本")

        # 添加第二个表格
        table2 = doc.add_table(rows=2, cols=2)
        table2.cell(0, 0).text = "表头1"
        table2.cell(0, 1).text = "表头2"

        doc_path = self.temp_path / filename
        doc.save(str(doc_path))
        return str(doc_path)

    def test_multiple_table_rules_together(self):
        """测试多个表格规则一起工作"""
        doc_path = self.create_complex_table_document()

        # 应用表格宽度规则
        context = RuleContext(doc_path)
        width_rule = TableWidthRule({
            "table_width_percent": 95,
            "auto_adjust_columns": True
        })
        width_result = width_rule.apply(context)
        self.assertTrue(width_result.success)

        # 应用表格边框规则
        border_rule = TableBorderRule({
            "border_size": 1,
            "border_color": "000000"
        })
        border_result = border_rule.apply(context)
        self.assertTrue(border_result.success)

        # 验证文档仍然有效
        doc = Document(doc_path)
        self.assertGreater(len(doc.tables), 0)

    def test_table_rules_with_engine(self):
        """测试表格规则在规则引擎中工作"""
        from core.engine import RuleEngine

        doc_path = self.create_complex_table_document()

        # 创建规则引擎
        engine = RuleEngine()
        engine._load_rules()

        # 获取表格规则
        rules_info = engine.get_rules_info()
        table_rules = [
            {"rule_id": "TableWidthRule", "params": {"table_width_percent": 95}},
            {"rule_id": "TableBorderRule", "params": {"border_size": 1}}
        ]

        # 执行规则
        result = engine.execute(doc_path, table_rules)

        # 验证结果
        self.assertEqual(result["status"], "success")
        self.assertEqual(len(result["results"]), 2)


if __name__ == '__main__':
    unittest.main()
