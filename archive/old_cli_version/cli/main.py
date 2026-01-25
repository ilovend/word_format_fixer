"""命令行接口主函数"""

import argparse
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

from ..core.fixer import RobustWordFixer
from ..core.config import load_config


def simple_fix_document(input_file, output_file=None, margin_left_cm=2.54, margin_right_cm=2.54, fix_numbering=True):
    """
    简单修复文档格式，主要调整页边距
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径，如为None则在原文件名后加_fixed
        margin_left_cm: 左边距（厘米）
        margin_right_cm: 右边距（厘米）
        fix_numbering: 是否修复编号列表
    """
    try:
        print(f"正在处理: {input_file}")
        
        # 加载文档
        doc = Document(input_file)
        
        # 设置页边距
        for section in doc.sections:
            # 确保页面大小
            if section.page_width is None:
                section.page_width = Cm(21.0)  # A4宽度
            if section.page_height is None:
                section.page_height = Cm(29.7)  # A4高度
            
            # 设置边距
            section.left_margin = Cm(margin_left_cm)
            section.right_margin = Cm(margin_right_cm)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            
            print(f"设置页边距:")
            print(f"  左边距: {margin_left_cm}cm")
            print(f"  右边距: {margin_right_cm}cm")
            print(f"  上边距: 2.54cm")
            print(f"  下边距: 2.54cm")
        
        # 设置字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 修复编号列表
        if fix_numbering:
            print("修复编号列表...")
            
            # 定义编号模式
            from ..core.utils import detect_numbering_patterns
            patterns = detect_numbering_patterns()
            
            fixed_count = 0
            for paragraph in doc.paragraphs:
                original_text = paragraph.text.strip()
                
                if not original_text or paragraph.style.name.startswith('Heading'):
                    continue
                
                # 检查各种编号模式
                matched = False
                for pattern_name, pattern in patterns.items():
                    match = re.match(pattern, original_text)
                    if match:
                        matched = True
                        # 提取编号和内容
                        number_part = match.group()
                        content = original_text[len(number_part):].strip()
                        
                        # 清除原有内容
                        paragraph.clear()
                        
                        # 添加内容
                        run = paragraph.add_run(content)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # 设置段落格式
                        if 'List Paragraph' in [style.name for style in doc.styles]:
                            paragraph.style = doc.styles['List Paragraph']
                        
                        # 设置缩进
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Cm(1.27)
                        paragraph_format.first_line_indent = Cm(-0.64)
                        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        paragraph_format.line_spacing = 1.5
                        
                        fixed_count += 1
                        break
            
            print(f"  修复了 {fixed_count} 个编号段落")
        
        # 调整表格
        if doc.tables:
            print(f"找到 {len(doc.tables)} 个表格，正在调整...")
            
            for table in doc.tables:
                # 设置表格居中对齐
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 自动调整列宽
                for col in table.columns:
                    col.width = Cm(4)  # 基础宽度
                
                # 添加边框
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        borders = ['top', 'left', 'bottom', 'right']
                        for border in borders:
                            border_elem = OxmlElement(f'w:{border}')
                            border_elem.set(qn('w:val'), 'single')
                            border_elem.set(qn('w:sz'), '4')
                            border_elem.set(qn('w:space'), '0')
                            border_elem.set(qn('w:color'), '000000')
                            tcPr.append(border_elem)
        
        # 保存文档
        if output_file is None:
            import os
            base, ext = os.path.splitext(input_file)
            output_file = f"{base}_fixed{ext}"
        
        doc.save(output_file)
        print(f"处理完成！文件保存到: {output_file}")
        
        return output_file
        
    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        return None


def main():
    """命令行主函数"""
    parser = argparse.ArgumentParser(description='Word文档格式修复工具')
    parser.add_argument('input', help='输入Word文档路径')
    parser.add_argument('-o', '--output', help='输出文件路径')
    
    # 字体设置
    parser.add_argument('--chinese-font', default='宋体',
                       help='中文字体，默认宋体')
    parser.add_argument('--western-font', default='Arial',
                       help='西文字体，默认Arial')
    parser.add_argument('--title-font', default='黑体',
                       help='标题字体，默认黑体')
    
    # 字号设置
    parser.add_argument('--font-size-body', type=int, default=12,
                       help='正文字号（磅），默认12')
    parser.add_argument('--font-size-title1', type=int, default=22,
                       help='一级标题字号（磅），默认22')
    parser.add_argument('--font-size-title2', type=int, default=18,
                       help='二级标题字号（磅），默认18')
    parser.add_argument('--font-size-title3', type=int, default=16,
                       help='三级标题字号（磅），默认16')
    
    # 页面设置
    parser.add_argument('--margin-top', type=float, default=2.54,
                       help='上边距（厘米），默认2.54cm')
    parser.add_argument('--margin-bottom', type=float, default=2.54,
                       help='下边距（厘米），默认2.54cm')
    parser.add_argument('--margin-left', type=float, default=2.54, 
                       help='左边距（厘米），默认2.54cm')
    parser.add_argument('--margin-right', type=float, default=2.54,
                       help='右边距（厘米），默认2.54cm')
    
    # 表格设置
    parser.add_argument('--table-width-percent', type=int, default=95,
                       help='表格宽度百分比，默认95%%')
    parser.add_argument('--auto-adjust-columns', action='store_true', default=True,
                       help='自动调整列宽，默认开启')
    parser.add_argument('--no-auto-adjust-columns', action='store_false', dest='auto_adjust_columns',
                       help='禁用自动调整列宽')
    
    # 模式设置
    parser.add_argument('--simple', action='store_true',
                       help='使用简单模式（只调整页边距和字体）')
    parser.add_argument('--numbering-only', action='store_true',
                       help='只修复编号列表')
    parser.add_argument('--config', help='配置文件路径')
    
    args = parser.parse_args()
    
    if args.numbering_only:
        # 只修复编号列表
        doc = Document(args.input)
        print("正在修复编号列表...")
        
        # 定义编号模式
        from ..core.utils import detect_numbering_patterns
        patterns = detect_numbering_patterns()
        
        fixed_count = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text.strip()
            
            if not original_text or paragraph.style.name.startswith('Heading'):
                continue
            
            # 检查各种编号模式
            matched = False
            for pattern_name, pattern in patterns.items():
                match = re.match(pattern, original_text)
                if match:
                    matched = True
                    # 提取编号和内容
                    number_part = match.group()
                    content = original_text[len(number_part):].strip()
                    
                    # 清除原有内容
                    paragraph.clear()
                    
                    # 添加内容
                    run = paragraph.add_run(content)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # 设置段落格式
                    if 'List Paragraph' in [style.name for style in doc.styles]:
                        paragraph.style = doc.styles['List Paragraph']
                    
                    # 设置缩进
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Cm(1.27)
                    paragraph_format.first_line_indent = Cm(-0.64)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph_format.line_spacing = 1.5
                    
                    fixed_count += 1
                    break
        
        print(f"  修复了 {fixed_count} 个编号段落")
        
        # 保存文档
        if args.output is None:
            import os
            base, ext = os.path.splitext(args.input)
            output_file = f"{base}_numbered{ext}"
        else:
            output_file = args.output
        
        doc.save(output_file)
        print(f"处理完成！文件保存到: {output_file}")
        return output_file
    
    elif args.simple:
        # 简单模式
        return simple_fix_document(
            args.input, 
            args.output, 
            args.margin_left, 
            args.margin_right
        )
    else:
        # 完整模式
        config = load_config(args.config) if args.config else {}
        config.update({
            # 字体设置
            'chinese_font': args.chinese_font,
            'western_font': args.western_font,
            'title_font': args.title_font,
            
            # 字号设置
            'font_size_body': args.font_size_body,
            'font_size_title1': args.font_size_title1,
            'font_size_title2': args.font_size_title2,
            'font_size_title3': args.font_size_title3,
            
            # 页面设置
            'page_margin_top_cm': args.margin_top,
            'page_margin_bottom_cm': args.margin_bottom,
            'page_margin_left_cm': args.margin_left,
            'page_margin_right_cm': args.margin_right,
            
            # 表格设置
            'table_width_percent': args.table_width_percent,
            'auto_adjust_columns': args.auto_adjust_columns,
        })
        
        fixer = RobustWordFixer(config)
        return fixer.fix_all(args.input, args.output)


if __name__ == "__main__":
    # 使用示例：
    # 1. 简单模式（只调整页边距）：
    #    python -m word_format_fixer.cli.main 文档.docx --simple --margin-left 2.0 --margin-right 2.0
    
    # 2. 完整模式：
    #    python -m word_format_fixer.cli.main 文档.docx -o 修复后.docx
    
    result = main()
    sys.exit(0 if result else 1)
