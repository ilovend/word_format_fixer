from docx import Document
from typing import Optional, Dict, Any

class RuleContext:
    """规则执行上下文"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.document: Optional[Document] = None
        self._load_document()
        self._cache: Dict[str, Any] = {}
        self.available_width_cm = 15.92  # 默认值，页面布局规则会更新它
        self.runtime_data = {}  # 用于规则间传递临时数据
        
        # 为了兼容测试用例，添加别名
        self.doc = self.document
        self.file_path = self.document_path
    
    def _load_document(self):
        """加载文档"""
        try:
            self.document = Document(self.document_path)
        except Exception as e:
            raise Exception(f"加载文档失败: {str(e)}")
    
    def get_document(self) -> Document:
        """获取文档对象"""
        if not self.document:
            self._load_document()
        return self.document
    
    def get_file_path(self) -> str:
        """获取文件路径"""
        return self.document_path
    
    def get_paragraphs(self) -> list:
        """获取文档中的所有段落"""
        return self.get_document().paragraphs
    
    def get_tables(self) -> list:
        """获取文档中的所有表格"""
        return self.get_document().tables
    
    def get_document_statistics(self) -> Dict[str, int]:
        """获取文档统计信息"""
        doc = self.get_document()
        return {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    
    def set_cache(self, key: str, value: Any):
        """设置缓存"""
        self._cache[key] = value
    
    def get_cache(self, key: str, default: Any = None) -> Any:
        """获取缓存"""
        return self._cache.get(key, default)
    
    def clear_cache(self):
        """清除缓存"""
        self._cache.clear()
    
    def save_document(self, output_path: Optional[str] = None):
        """保存文档"""
        save_path = output_path or self.document_path
        if self.document:
            self.document.save(save_path)
            return True
        return False