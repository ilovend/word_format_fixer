from docx import Document
from typing import Optional, List, Any

class DocumentWrapper:
    """文档包装器，提供统一的文档操作接口"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self._document: Optional[Document] = None
        self._load_document()
    
    def _load_document(self):
        """加载文档"""
        try:
            self._document = Document(self.document_path)
        except Exception as e:
            raise Exception(f"加载文档失败: {str(e)}")
    
    def get_document(self) -> Document:
        """获取原始文档对象"""
        if not self._document:
            self._load_document()
        return self._document
    
    def save(self, output_path: Optional[str] = None):
        """保存文档"""
        save_path = output_path or self.document_path
        if self._document:
            self._document.save(save_path)
            return True
        return False
    
    def get_paragraphs(self):
        """获取所有段落"""
        if self._document:
            return self._document.paragraphs
        return []
    
    def get_tables(self):
        """获取所有表格"""
        if self._document:
            return self._document.tables
        return []
    
    def get_styles(self):
        """获取所有样式"""
        if self._document:
            return self._document.styles
        return []
    
    def get_sections(self):
        """获取所有节"""
        if self._document:
            return self._document.sections
        return []