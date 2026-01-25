# -*- coding: utf-8 -*-
"""字体标准化规则集"""

from rules.base_rule import BaseRule, RuleResult
from docx.shared import Pt
from docx.oxml.ns import qn
from schemas.rule_params import (
    RuleConfigSchema, 
    FontParam, 
    SizeParam,
    RangeParam
)


class FontNameRule(BaseRule):
    """字体名称标准化规则 - 设置中英文字体"""

    display_name = "字体名称标准化"
    category = "字体规则"
    description = "为文档设置统一的中文字体和西文字体"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        FontParam(
            name="chinese_font",
            display_name="中文字体",
            default="宋体",
            description="中文内容使用的字体"
        ),
        FontParam(
            name="western_font",
            display_name="西文字体",
            default="Arial",
            description="英文和数字使用的字体"
        ),
    ])

    def apply(self, doc_context) -> RuleResult:
        """
        应用字体名称规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.config['western_font']
                run._element.rPr.rFonts.set(
                    qn('w:eastAsia'),
                    self.config['chinese_font']
                )
                fixed_count += 1

        # 处理表格中的文本
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.config['western_font']
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'),
                                self.config['chinese_font']
                            )
                            fixed_count += 1

        details.append(f"中文字体: {self.config['chinese_font']}, 西文字体: {self.config['western_font']}")
        details.append(f"标准化了 {fixed_count} 个文本运行的字体")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )


class TitleFontRule(BaseRule):
    """标题字体规则 - 标题使用专用字体"""

    display_name = "标题字体设置"
    category = "字体规则"
    description = "为文档标题设置专用字体"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        FontParam(
            name="title_font",
            display_name="标题中文字体",
            default="黑体",
            description="标题使用的中文字体"
        ),
    ])

    def apply(self, doc_context) -> RuleResult:
        """
        应用标题字体规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                for run in paragraph.runs:
                    run.font.name = 'Arial'  # 西文字体固定为Arial
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'),
                        self.config['title_font']
                    )
                    fixed_count += 1

        details.append(f"标题使用字体: {self.config['title_font']}")
        details.append(f"标准化了 {fixed_count} 个标题文本运行的字体")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )


class FontSizeRule(BaseRule):
    """字号标准化规则 - 设置标题和正文的字号"""

    display_name = "字号标准化"
    category = "字体规则"
    description = "统一设置文档中正文和各级标题的字号"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        SizeParam(
            name="font_size_body",
            display_name="正文字号",
            default=12,
            min_value=8,
            max_value=36,
            step=0.5,
            unit="pt",
            description="正文内容的字号大小"
        ),
        SizeParam(
            name="font_size_title1",
            display_name="一级标题字号",
            default=22,
            min_value=12,
            max_value=48,
            step=1,
            unit="pt",
            description="一级标题（H1）的字号"
        ),
        SizeParam(
            name="font_size_title2",
            display_name="二级标题字号",
            default=18,
            min_value=10,
            max_value=36,
            step=1,
            unit="pt",
            description="二级标题（H2）的字号"
        ),
        SizeParam(
            name="font_size_title3",
            display_name="三级标题字号",
            default=16,
            min_value=10,
            max_value=30,
            step=1,
            unit="pt",
            description="三级标题（H3）的字号"
        ),
        RangeParam(
            name="min_font_size",
            display_name="最小允许字号",
            default=10,
            min_value=6,
            max_value=16,
            step=1,
            unit="pt",
            description="小于此字号的文本会被自动调整为正文字号"
        ),
    ])

    def apply(self, doc_context) -> RuleResult:
        """
        应用字号规则
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                # 根据标题级别设置字号
                for run in paragraph.runs:
                    if paragraph.style.name == 'Heading 1':
                        run.font.size = Pt(self.config['font_size_title1'])
                    elif paragraph.style.name == 'Heading 2':
                        run.font.size = Pt(self.config['font_size_title2'])
                    elif paragraph.style.name == 'Heading 3':
                        run.font.size = Pt(self.config['font_size_title3'])
                    else:
                        run.font.size = Pt(self.config['font_size_body'])
                    fixed_count += 1
            else:
                # 正文：如果字号太小则标准化
                for run in paragraph.runs:
                    if run.font.size is None or run.font.size.pt < self.config['min_font_size']:
                        run.font.size = Pt(self.config['font_size_body'])
                        fixed_count += 1

        details.append(f"正文字号: {self.config['font_size_body']}pt")
        details.append(f"一级标题字号: {self.config['font_size_title1']}pt")
        details.append(f"二级标题字号: {self.config['font_size_title2']}pt")
        details.append(f"三级标题字号: {self.config['font_size_title3']}pt")
        details.append(f"标准化了 {fixed_count} 个文本运行的字号")

        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
