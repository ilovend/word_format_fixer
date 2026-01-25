# -*- coding: utf-8 -*-
"""表格宽度规则"""

from rules.base_rule import BaseRule, RuleResult
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
import re
from schemas.rule_params import (
    RuleConfigSchema,
    RangeParam,
    BoolParam,
    EnumParam
)


class TableWidthRule(BaseRule):
    """表格宽度规则 - 优化表格宽度，支持复杂表格结构"""

    display_name = "表格宽度优化"
    category = "表格规则"
    description = "自动调整表格宽度和列宽，支持合并单元格和嵌套表格"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        RangeParam(
            name="table_width_percent",
            display_name="表格宽度百分比",
            default=95,
            min_value=50,
            max_value=100,
            step=5,
            unit="%",
            description="表格宽度占页面可用宽度的百分比"
        ),
        EnumParam(
            name="table_alignment",
            display_name="表格对齐方式",
            options=[
                {"value": "center", "label": "居中"},
                {"value": "left", "label": "左对齐"},
                {"value": "right", "label": "右对齐"},
            ],
            default="center",
            description="表格在页面中的对齐方式"
        ),
        BoolParam(
            name="auto_adjust_columns",
            display_name="自动调整列宽",
            default=True,
            description="根据内容自动计算最优列宽"
        ),
    ])
    
    # 对齐方式映射
    ALIGN_MAP = {
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'right': WD_TABLE_ALIGNMENT.RIGHT,
    }
    
    def apply(self, doc_context) -> RuleResult:
        """
        核心执行逻辑
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        if not document.tables:
            details.append("文档中没有表格，跳过表格宽度优化")
            return RuleResult(
                rule_id=self.rule_id,
                success=True,
                fixed_count=fixed_count,
                details=details
            )
        
        details.append(f"开始优化表格（共 {len(document.tables)} 个）...")

        # 获取第一个分区的页面信息
        section = document.sections[0]
        available_width_cm = (section.page_width.cm -
                             section.left_margin.cm -
                             section.right_margin.cm)

        for table_idx, table in enumerate(document.tables):
            # 检查是否有合并单元格或嵌套表格
            has_merged_cells = self._check_merged_cells(table)
            has_nested_tables = self._check_nested_tables(table)

            # 计算表格宽度
            table_width_cm = available_width_cm * self.config['table_width_percent'] / 100
            table.width = Cm(table_width_cm)
            
            # 设置对齐方式
            alignment = self.config.get('table_alignment', 'center')
            table.alignment = self.ALIGN_MAP.get(alignment, WD_TABLE_ALIGNMENT.CENTER)
            
            # 处理列宽
            col_count = len(table.columns)
            
            if self.config['auto_adjust_columns']:
                # 自动调整列宽
                self._auto_adjust_column_widths(table, table_width_cm, details)
            else:
                # 平均分配列宽
                col_width_cm = table_width_cm / col_count
                for col in table.columns:
                    col.width = Cm(col_width_cm)
            
            # 处理嵌套表格
            if has_nested_tables:
                self._process_nested_tables(table, document, details)
            
            fixed_count += 1
        
        details.append(f"总共优化了 {fixed_count} 个表格的宽度")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
    
    def _check_merged_cells(self, table):
        """检查表格是否有合并单元格"""
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                for child in tcPr.iter():
                    local_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if local_name in ['vMerge', 'hMerge']:
                        return True
        return False
    
    def _check_nested_tables(self, table):
        """检查表格是否有嵌套表格"""
        for row in table.rows:
            for cell in row.cells:
                if len(cell.tables) > 0:
                    return True
        return False
    
    def _process_nested_tables(self, table, document, details):
        """处理嵌套表格"""
        section = document.sections[0]
        available_width_cm = (section.page_width.cm -
                             section.left_margin.cm -
                             section.right_margin.cm)

        for row in table.rows:
            for cell in row.cells:
                for nested_table in cell.tables:
                    nested_table_width_cm = cell.width.cm * 0.9 if cell.width else available_width_cm * 0.7
                    nested_table.width = Cm(nested_table_width_cm)
    
    def _auto_adjust_column_widths(self, table, table_width_cm, details):
        """自动调整列宽 - 支持复杂表格"""
        col_count = len(table.columns)
        max_lengths = [0] * col_count
        
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                gridSpan = None
                for child in tcPr.iter():
                    local_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if local_name == 'gridSpan':
                        gridSpan = child
                        break
                span = 1
                if gridSpan is not None:
                    val = None
                    for attr in gridSpan.attrib:
                        local_attr_name = attr.split('}')[-1] if '}' in attr else attr
                        if local_attr_name == 'val':
                            val = gridSpan.attrib[attr]
                            break
                    if val:
                        span = int(val)
                
                text = cell.text.strip()
                if text:
                    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
                    english_chars = len(text) - chinese_chars
                    length = (english_chars + chinese_chars * 2) / span
                    
                    for k in range(j, min(j + span, col_count)):
                        max_lengths[k] = max(max_lengths[k], length)
        
        total_length = sum(max_lengths) if sum(max_lengths) > 0 else 1
        
        for j, col in enumerate(table.columns):
            width_ratio = max_lengths[j] / total_length
            col_width_cm = table_width_cm * width_ratio
            col.width = Cm(col_width_cm)
