# -*- coding: utf-8 -*-
"""表格边框规则"""

from rules.base_rule import BaseRule, RuleResult
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Cm
from schemas.rule_params import (
    RuleConfigSchema,
    FontParam,
    SizeParam,
    ColorParam,
    RangeParam,
    BoolParam
)


class TableBorderRule(BaseRule):
    """表格边框规则 - 为表格添加边框并格式化单元格"""

    display_name = "表格边框和格式"
    category = "表格规则"
    description = "为表格添加统一边框样式，并格式化表头和内容单元格"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        RangeParam(
            name="border_size",
            display_name="边框粗细",
            default=4,
            min_value=1,
            max_value=12,
            step=1,
            unit="",
            description="表格边框的粗细（单位：1/8磅）"
        ),
        ColorParam(
            name="border_color",
            display_name="边框颜色",
            default="#000000",
            description="表格边框的颜色"
        ),
        BoolParam(
            name="add_table_header_format",
            display_name="格式化表头",
            default=True,
            description="是否为第一行（表头）设置特殊格式"
        ),
        ColorParam(
            name="table_header_bg_color",
            display_name="表头背景色",
            default="#E3E3E3",
            description="表头行的背景颜色"
        ),
        SizeParam(
            name="font_size_table_header",
            display_name="表头字号",
            default=14,
            min_value=10,
            max_value=24,
            step=1,
            unit="pt",
            description="表头文字的字号"
        ),
        SizeParam(
            name="font_size_table_content",
            display_name="内容字号",
            default=12,
            min_value=8,
            max_value=20,
            step=1,
            unit="pt",
            description="表格内容文字的字号"
        ),
        FontParam(
            name="chinese_font",
            display_name="中文字体",
            default="宋体",
            description="表格中文文字使用的字体"
        ),
        FontParam(
            name="western_font",
            display_name="西文字体",
            default="Arial",
            description="表格西文文字使用的字体"
        ),
    ])
    
    def apply(self, doc_context) -> RuleResult:
        """
        核心执行逻辑
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        if not document.tables:
            details.append("文档中没有表格，跳过表格边框和格式设置")
            return RuleResult(
                rule_id=self.rule_id,
                success=True,
                fixed_count=fixed_count,
                details=details
            )
        
        details.append(f"开始为表格添加边框和格式（共 {len(document.tables)} 个）...")
        
        border_size = self.config.get('border_size', 4)
        border_color = self._parse_color_hex(self.config.get('border_color', '#000000'))
        
        for table_idx, table in enumerate(document.tables):
            # 为表格添加边框
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell, border_size, border_color)
                    fixed_count += 1
            
            # 格式化表格单元格
            self._format_table_cells(table)
        
        details.append(f"总共处理了 {fixed_count} 个表格单元格")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
    
    def _parse_color_hex(self, color_value):
        """解析颜色值为十六进制字符串"""
        if isinstance(color_value, str):
            if color_value.startswith('#'):
                return color_value[1:]
            return color_value
        return "000000"
    
    def _set_cell_border(self, cell, border_size: int = 4, border_color: str = "000000"):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        borders = ['top', 'left', 'bottom', 'right']
        for border in borders:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), str(border_size))
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), border_color)
            tcPr.append(border_elem)
    
    def _format_table_cells(self, table):
        """格式化表格单元格"""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 设置单元格垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 设置字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.config['western_font']
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'), 
                            self.config['chinese_font']
                        )
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # 根据位置设置字号
                        if i == 0 and self.config['add_table_header_format']:
                            run.font.size = Pt(self.config['font_size_table_header'])
                            run.font.bold = True
                        else:
                            run.font.size = Pt(self.config['font_size_table_content'])
                
                # 设置表头背景色
                if i == 0 and self.config['add_table_header_format']:
                    bg_color = self._parse_color_hex(self.config.get('table_header_bg_color', '#E3E3E3'))
                    self._set_cell_background(cell, bg_color)
                
                # 设置单元格边距
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.left_indent = Cm(0.2)
                    cell.paragraphs[0].paragraph_format.right_indent = Cm(0.2)
    
    def _set_cell_background(self, cell, color):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._element.tcPr.append(shading)
