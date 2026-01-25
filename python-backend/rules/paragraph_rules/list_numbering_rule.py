# -*- coding: utf-8 -*-
"""编号列表规则"""

import re
from rules.base_rule import BaseRule, RuleResult
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from schemas.rule_params import (
    RuleConfigSchema,
    FontParam,
    SizeParam,
    ColorParam,
    RangeParam
)


class ListNumberingRule(BaseRule):
    """编号列表规则 - 修复文档中的编号列表和项目符号"""

    display_name = "编号列表标准化"
    category = "段落规则"
    description = "修复和标准化文档中的编号列表和项目符号格式"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        FontParam(
            name="chinese_font",
            display_name="中文字体",
            default="宋体",
            description="列表项使用的中文字体"
        ),
        FontParam(
            name="western_font",
            display_name="西文字体",
            default="Arial",
            description="列表项使用的西文字体"
        ),
        SizeParam(
            name="font_size_body",
            display_name="列表字号",
            default=12,
            min_value=8,
            max_value=24,
            step=0.5,
            unit="pt",
            description="列表项文本的字号"
        ),
        ColorParam(
            name="text_color",
            display_name="文本颜色",
            default="#000000",
            description="列表项文本的颜色"
        ),
        RangeParam(
            name="list_indent",
            display_name="列表缩进",
            default=1.27,
            min_value=0.5,
            max_value=3.0,
            step=0.1,
            unit="cm",
            description="列表项的左缩进距离"
        ),
        RangeParam(
            name="line_spacing",
            display_name="行间距",
            default=1.5,
            min_value=1.0,
            max_value=3.0,
            step=0.1,
            unit="倍",
            description="列表项的行距倍数"
        ),
    ])
    
    def apply(self, doc_context) -> RuleResult:
        """
        核心执行逻辑
        :param doc_context: 文档上下文对象
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        # 获取编号模式
        patterns = self.detect_numbering_patterns()
        
        for paragraph in document.paragraphs:
            original_text = paragraph.text.strip()
            
            if not original_text or paragraph.style.name.startswith('Heading'):
                continue
            
            # 处理不需要的项目符号（如·）
            bullet_patterns = [
                r'^\s*·\s+',  # 中文点号
                r'^\s*\*\s+',  # 星号
                r'^\s*-\s+',   # 连字符
                r'^\s+•\s+',   # 实心圆点
            ]
            
            bullet_matched = False
            for bullet_pattern in bullet_patterns:
                bullet_match = re.match(bullet_pattern, original_text)
                if bullet_match:
                    bullet_matched = True
                    # 提取内容（去掉项目符号）
                    content = original_text[len(bullet_match.group()):].strip()
                    
                    # 清除原有内容
                    paragraph.clear()
                    
                    # 添加内容
                    run = paragraph.add_run(content)
                    run.font.name = self.config['western_font']
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 
                        self.config['chinese_font']
                    )
                    run.font.color.rgb = self._parse_color(self.config.get('text_color', '#000000'))
                    run.font.size = Pt(self.config['font_size_body'])
                    
                    # 设置为无序列表样式
                    if 'List Paragraph' in [style.name for style in document.styles]:
                        paragraph.style = document.styles['List Paragraph']
                    
                    # 设置缩进
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Cm(self.config.get('list_indent', 1.27))
                    paragraph_format.first_line_indent = Cm(-0.64)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph_format.line_spacing = self.config.get('line_spacing', 1.5)
                    
                    fixed_count += 1
                    break
            
            if bullet_matched:
                continue
            
            # 检查各种编号模式
            for pattern_name, pattern in patterns.items():
                match = re.match(pattern, original_text)
                if match:
                    self._format_numbered_paragraph(paragraph, pattern_name, match, original_text, document)
                    fixed_count += 1
                    break
        
        details.append(f"总共修复了 {fixed_count} 个编号或项目符号段落")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
    
    def detect_numbering_patterns(self):
        """检测文档中的编号模式"""
        patterns = {
            'arabic_with_dot': r'^\s*(\d+)\.\s+',  # 1. 项目
            'arabic_with_paren': r'^\s*(\d+)\)\s+',  # 1) 项目
            'chinese_with_dot': r'^\s*([一二三四五六七八九十]+)\.\s+',  # 一. 项目
            'chinese_with_bracket': r'^\s*([一二三四五六七八九十]+)、\s+',  # 一、项目
            'chinese_with_paren': r'^\s*\(([一二三四五六七八九十]+)\)\s+',  # (一) 项目
            'lower_alpha_with_dot': r'^\s*([a-z])\.\s+',  # a. 项目
            'upper_alpha_with_dot': r'^\s*([A-Z])\.\s+',  # A. 项目
            'lower_roman_with_dot': r'^\s*([ivxlcdm]+)\.\s+',  # i. 项目
            'upper_roman_with_dot': r'^\s*([IVXLCDM]+)\.\s+',  # I. 项目
        }
        return patterns
    
    def _format_numbered_paragraph(self, paragraph, pattern_type: str, match, original_text: str, document):
        """格式化编号段落"""
        # 提取编号和内容
        number_part = match.group()
        content = original_text[len(number_part):].strip()
        
        # 清除原有内容
        paragraph.clear()
        
        # 添加内容
        run = paragraph.add_run(content)
        run.font.name = self.config['western_font']
        run._element.rPr.rFonts.set(
            qn('w:eastAsia'), 
            self.config['chinese_font']
        )
        run.font.color.rgb = self._parse_color(self.config.get('text_color', '#000000'))
        run.font.size = Pt(self.config['font_size_body'])
        
        # 设置段落格式
        if 'List Paragraph' in [style.name for style in document.styles]:
            paragraph.style = document.styles['List Paragraph']
        else:
            paragraph.style = document.styles['Normal']
        
        # 设置缩进
        paragraph_format = paragraph.paragraph_format
        list_indent = self.config.get('list_indent', 1.27)
        
        # 根据编号类型设置不同缩进
        if pattern_type.startswith('arabic') or pattern_type.startswith('chinese'):
            # 主要列表项
            paragraph_format.left_indent = Cm(list_indent)
            paragraph_format.first_line_indent = Cm(-0.64)
        elif pattern_type.startswith('lower_') or pattern_type.startswith('upper_'):
            # 次级列表项
            paragraph_format.left_indent = Cm(list_indent * 2)
            paragraph_format.first_line_indent = Cm(-0.64)
        else:
            # 默认缩进
            paragraph_format.left_indent = Cm(list_indent)
            paragraph_format.first_line_indent = Cm(-0.64)
        
        # 设置行距
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = self.config.get('line_spacing', 1.5)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
    
    def _parse_color(self, color_value):
        """解析颜色值"""
        from docx.shared import RGBColor
        
        # 如果是字符串（十六进制）
        if isinstance(color_value, str):
            if color_value.startswith('#'):
                color_value = color_value[1:]
            if len(color_value) == 6:
                r = int(color_value[0:2], 16)
                g = int(color_value[2:4], 16)
                b = int(color_value[4:6], 16)
                return RGBColor(r, g, b)
        # 如果是元组
        elif isinstance(color_value, (tuple, list)) and len(color_value) == 3:
            return RGBColor(*color_value)
        
        # 默认黑色
        return RGBColor(0, 0, 0)
