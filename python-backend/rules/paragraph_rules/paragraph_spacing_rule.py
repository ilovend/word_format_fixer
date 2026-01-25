# -*- coding: utf-8 -*-
"""段落间距统一规则"""

from rules.base_rule import BaseRule, RuleResult
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING
from schemas.rule_params import RuleConfigSchema, RangeParam


class ParagraphSpacingRule(BaseRule):
    """段落间距统一规则 - 统一文档中所有段落的间距和缩进"""

    display_name = "段落间距统一"
    category = "段落规则"
    description = "统一设置文档中正文和表格内段落的间距、缩进和行距"
    
    # 参数 Schema 定义
    param_schema = RuleConfigSchema(params=[
        RangeParam(
            name="body_left_indent",
            display_name="正文左缩进",
            default=0,
            min_value=0,
            max_value=5,
            step=0.1,
            unit="cm",
            description="正文段落左侧缩进距离"
        ),
        RangeParam(
            name="body_right_indent",
            display_name="正文右缩进",
            default=0,
            min_value=0,
            max_value=5,
            step=0.1,
            unit="cm",
            description="正文段落右侧缩进距离"
        ),
        RangeParam(
            name="body_space_before",
            display_name="段前间距",
            default=0,
            min_value=0,
            max_value=2,
            step=0.1,
            unit="cm",
            description="正文段落与上一段落的间距"
        ),
        RangeParam(
            name="body_space_after",
            display_name="段后间距",
            default=0.33,
            min_value=0,
            max_value=2,
            step=0.1,
            unit="cm",
            description="正文段落与下一段落的间距"
        ),
        RangeParam(
            name="body_line_spacing",
            display_name="行间距",
            default=1.5,
            min_value=1.0,
            max_value=3.0,
            step=0.1,
            unit="倍",
            description="正文段落的行距倍数"
        ),
        RangeParam(
            name="table_left_indent",
            display_name="表格内左缩进",
            default=0.2,
            min_value=0,
            max_value=2,
            step=0.1,
            unit="cm",
            description="表格内段落左侧缩进"
        ),
        RangeParam(
            name="table_right_indent",
            display_name="表格内右缩进",
            default=0.2,
            min_value=0,
            max_value=2,
            step=0.1,
            unit="cm",
            description="表格内段落右侧缩进"
        ),
    ])
    
    def apply(self, doc_context):
        """
        应用段落间距规则
        :param doc_context: 文档上下文
        :return: 规则执行结果
        """
        document = doc_context.get_document()
        fixed_count = 0
        details = []
        
        # 处理普通段落
        for paragraph in document.paragraphs:
            # 跳过标题段落
            if not paragraph.style.name.startswith('Heading'):
                # 设置段落格式
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Cm(self.config['body_left_indent'])
                paragraph_format.right_indent = Cm(self.config['body_right_indent'])
                paragraph_format.space_before = Cm(self.config['body_space_before'])
                paragraph_format.space_after = Cm(self.config['body_space_after'])
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph_format.line_spacing = self.config['body_line_spacing']
                fixed_count += 1
        
        # 处理表格中的段落
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Cm(self.config['table_left_indent'])
                        paragraph_format.right_indent = Cm(self.config['table_right_indent'])
                        fixed_count += 1
        
        if fixed_count > 0:
            details.append(f"统一了 {fixed_count} 个段落的间距和缩进")
            details.append(f"行间距: {self.config['body_line_spacing']}倍")
        else:
            details.append("文档中没有需要处理的段落")
        
        return RuleResult(
            rule_id=self.rule_id,
            success=True,
            fixed_count=fixed_count,
            details=details
        )
